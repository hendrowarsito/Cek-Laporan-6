"""
CekLaporan v7.0 — KJPP SRR
Powered by Claude AI · Created by HW

Peningkatan dari v6:
- Upload & parsing lembar kerja XLSX (komparasi angka laporan vs workbook)
- Artifact detector: deteksi teks copy-paste dari laporan lain
- Placeholder detector: deteksi nilai "xx", tabel kosong, dll
- Numbering checker: deteksi penomoran ganda tabel/gambar/sub-bab
- Enhanced prompts: lebih spesifik dan mendalam untuk tiap mode
- Download laporan review sebagai teks
- pdfplumber untuk ekstraksi PDF lebih akurat
"""

import os
import re
import json
import time
import io
import anthropic
import streamlit as st
from docx import Document
from datetime import datetime

try:
    import pdfplumber
    PDF_ENGINE = "pdfplumber"
except ImportError:
    import PyPDF2
    PDF_ENGINE = "PyPDF2"

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    import gspread
    from google.oauth2.service_account import Credentials
    GSPREAD_AVAILABLE = True
except ImportError:
    GSPREAD_AVAILABLE = False

# ──────────────────────────────────────────────
# KONSTANTA
# ──────────────────────────────────────────────
MODEL      = "claude-sonnet-4-5"
MAX_TOKENS = 8192
MAX_CHARS  = 45000   # dinaikkan dari 40000

SHEET_RIWAYAT   = "riwayat_audit"
SHEET_REFERENSI = "data_laporan"

GSPREAD_SCOPES = [
    "https://spreadsheets.google.com/feeds",
    "https://www.googleapis.com/auth/drive",
]

# Kata-kata yang mengindikasikan artefak copy-paste dari laporan lain
ARTIFACT_KEYWORDS = [
    # Nama entitas yang tidak seharusnya ada
    "AMP", "SBPL", "MDK", "MDR", "TPIA",
    "Petrosea Services Solutions",
    "Singaraja Putra",  # hanya jika bukan obyek
    # Frasa yang biasanya menandakan copy-paste
    "pasar Singapura",
    "peleburan nikel",
    "nikel dan barang",
]

PLACEHOLDER_PATTERNS = [
    r'\bRp\s+xx\b',
    r'\bUSD?\s+xx\b',
    r'\bRp\s*\.\s*\.\s*\.',
    r'\bxxxx\b',
    r'\b\[.*?diisi.*?\]',
    r'\bTABEL\s*$',          # Tabel kosong
    r'\bGambar\s*$',
    r'(?<!\w)xx(?!\w)',
    r'\bTBD\b',
    r'\bN/A\b(?!\s*[\d,])',  # N/A bukan angka
    r'\.{5,}',               # banyak titik beruntun (placeholder)
]

SEVERITY_CONFIG = {
    "kritis":   {"emoji": "🔴", "color": "#c0392b", "bg": "#fff0f0"},
    "kritikal": {"emoji": "🔴", "color": "#c0392b", "bg": "#fff0f0"},  # alias backward compat
    "mayor":    {"emoji": "🟠", "color": "#e67e22", "bg": "#fff5ec"},
    "minor":    {"emoji": "🟡", "color": "#d4860a", "bg": "#fff8e6"},
    "ok":       {"emoji": "🟢", "color": "#1a9e67", "bg": "#edfaf4"},
    "info":     {"emoji": "🔵", "color": "#1e6fbf", "bg": "#eef4ff"},
    "mismatch": {"emoji": "🔶", "color": "#7c3aed", "bg": "#f5f3ff"},
}

# Konfigurasi seksi tampilan temuan — sesuai CLAUDE.md Bagian A-J
SECTION_CONFIG = [
    {
        "num": "A", "key": "angka",
        "title": "VERIFIKASI ANGKA KUNCI (Laporan vs LK)",
        "categories": ["Konsistensi Nilai", "Formula Resume", "Konsistensi Tanggal",
                       "Konsistensi Alamat", "Konsistensi Luas"],
        "style": "table_angka",  # No | Lokasi | Parameter | Nilai LK | Nilai LP | Status
    },
    {
        "num": "B", "key": "artefak",
        "title": "TEKS ASING / SISA COPY-PASTE",
        "categories": ["Artefak Copy-Paste"],
        "style": "narrative",
    },
    {
        "num": "C", "key": "placeholder",
        "title": "PLACEHOLDER / DATA BELUM DIISI",
        "categories": ["Placeholder", "Placeholder Belum Diisi"],
        "style": "table_lokasi_keterangan",
    },
    {
        "num": "D", "key": "penomoran",
        "title": "INKONSISTENSI PENOMORAN",
        "categories": ["Penomoran", "Penomoran Ganda", "Penomoran Sub-Bab"],
        "style": "table_lokasi_masalah",
    },
    {
        "num": "E", "key": "nama",
        "title": "INKONSISTENSI NAMA ENTITAS",
        "categories": ["Nama Perusahaan"],
        "style": "table_nama",
    },
    {
        "num": "F", "key": "ejaan",
        "title": "INKONSISTENSI EJAAN & PENULISAN",
        "categories": ["Ejaan & Penulisan", "Inkonsistensi Ejaan"],
        "style": "table_ejaan",
    },
    {
        "num": "G", "key": "typo",
        "title": "KESALAHAN KETIK & BAHASA",
        "categories": ["Typo"],
        "style": "table_typo",
    },
    {
        "num": "H", "key": "logis",
        "title": "INKONSISTENSI LOGIS / SUBSTANSI",
        "categories": ["Logis/Substansi"],
        "style": "table_lokasi_masalah",
    },
    {
        "num": "I", "key": "daftar",
        "title": "DAFTAR ISI, DAFTAR GAMBAR, DAFTAR TABEL",
        "categories": ["Daftar Isi", "Daftar Gambar", "Daftar Tabel"],
        "style": "table_lokasi_masalah",
    },
    {
        "num": "J", "key": "standar",
        "title": "KEPATUHAN STANDAR SPI / KEPI / POJK",
        "categories": ["KEPI/SPI/POJK", "SPI", "POJK"],
        "style": "table_lokasi_masalah",
    },
    {
        "num": "—", "key": "lainlain",
        "title": "TEMUAN LAINNYA",
        "categories": [],  # catch-all
        "style": "table_lokasi_masalah",
    },
]

# ──────────────────────────────────────────────
# MODE CONFIG — instruksi untuk Claude
# ──────────────────────────────────────────────
MODE_CONFIG = {
    "🔍 Pre-Check": {
        "key": "precheck",
        "desc": "Orientasi cepat — temukan masalah KRITIS dalam 60 detik.",
        "instruction": (
            "Lakukan Pre-Check cepat. Bangun model mental laporan ini terlebih dahulu:\n"
            "- Apa nama objek/perusahaan yang dinilai?\n"
            "- Siapa kliennya? Di mana lokasinya?\n"
            "- Berapa nilai indikasinya?\n\n"
            "Kemudian fokus pada TIGA hal paling kritikal:\n"
            "1. ARTEFAK COPY-PASTE: Ada nama entitas/singkatan yang tidak pernah diperkenalkan "
            "dalam laporan ini? Cek setiap nama perusahaan, singkatan, dan konteks narasi. "
            "Laporan saham sangat rentan — setiap nama asing = temuan KRITIS.\n"
            "2. PLACEHOLDER BELUM DIISI: 'Rp xx', 'USD xx', '[nilai]', tabel berisi "
            "hanya kata 'Tabel' tanpa data, bagian yang tampak belum selesai.\n"
            "3. FOOTER/HEADER SALAH: Baca footer setiap halaman — apakah menyebut "
            "nama proyek yang benar?\n\n"
            "Ringkas dalam 5-10 temuan prioritas tinggi. Setiap temuan harus spesifik: "
            "sebutkan halaman, kutip teks bermasalah."
        ),
    },
    "🧠 Deep Audit": {
        "key": "deepaudit",
        "desc": "Audit editor senior — baca seluruh laporan dengan tiga lapisan pemeriksaan.",
        "instruction": (
            "Lakukan DEEP AUDIT sebagai editor senior. Baca laporan dari halaman pertama "
            "hingga terakhir. Jalankan TIGA lapisan pemeriksaan secara simultan:\n\n"

            "=== LAPISAN 1: LINGUISTIK (Kata demi Kata) ===\n"
            "Deteksi secara MANDIRI (bukan hanya dari daftar pola):\n"
            "• Setiap kata yang terasa tidak lengkap atau ada huruf lebih\n"
            "• Kata yang terulang dua kali berturutan ('terdiri terdiri dari')\n"
            "• Huruf visual mirip tapi beda: l vs I, 0 vs O, rn vs m\n"
            "• Tanda baca menggantikan huruf ('era!' → 'erat', '1nflasi' → 'Inflasi')\n"
            "• Spasi di tengah kata ('tu run' → 'turun')\n"
            "• Awal kalimat tidak kapital, kalimat tidak berakhir tanda titik\n"
            "• Format angka salah: 'Rp 283.118, juta' (koma di posisi salah)\n"
            "• Ejaan KBBI: 'batu bara' (dua kata), 'bertanggung jawab', 'kerja sama'\n\n"

            "=== LAPISAN 2: KONTEKS & RELEVANSI (Paragraf demi Paragraf) ===\n"
            "• Catat semua nama entitas dan singkatan yang muncul\n"
            "• Tandai yang tidak pernah diperkenalkan → kemungkinan artefak copy-paste\n"
            "• Apakah konteks narasi relevan dengan bisnis yang dinilai?\n"
            "• Baca footer setiap halaman — apakah nama proyek benar?\n"
            "• Cari placeholder: 'Rp xx', tabel kosong, bagian belum selesai\n\n"

            "=== LAPISAN 3: KONSISTENSI (Dokumen Keseluruhan) ===\n"
            "• Angka yang sama — selalu sama di setiap bagian?\n"
            "• Nilai kesimpulan = nilai ringkasan eksekutif = nilai surat pengantar?\n"
            "• Nama perusahaan/objek — ejaan sama persis?\n"
            "• Nomor sub-bab urut dan logis?\n"
            "• Judul sub-bab mencerminkan isinya?\n"
            "• Cover dan footer menyebut jenis laporan yang sama?\n"
            "• Kronologis: tanggal inspeksi ≤ tanggal cut-off ≤ tanggal laporan?\n\n"

            "Laporkan SETIAP temuan dengan: halaman/bab spesifik, kutipan teks bermasalah, "
            "dan saran perbaikan konkret."
        ),
    },
    "📐 KEPI/MAPPI": {
        "key": "mappi",
        "desc": "Kepatuhan standar SPI, KEPI, MAPPI.",
        "instruction": (
            "Lakukan pemeriksaan kepatuhan standar KEPI/MAPPI dan SPI (Standar Penilaian Indonesia).\n\n"
            "ELEMEN WAJIB SPI yang harus ada:\n"
            "□ Pernyataan kepatuhan terhadap SPI\n"
            "□ Asumsi dan syarat pembatas\n"
            "□ Kualifikasi dan nomor izin penilai (MAPPI)\n"
            "□ Pembatasan penggunaan laporan\n"
            "□ Pernyataan independensi penilai\n"
            "□ Tanggal inspeksi / kunjungan lapangan\n"
            "□ Tujuan penilaian\n"
            "□ Dasar nilai (Nilai Pasar / Nilai Likuidasi / dll.)\n"
            "□ Pendekatan dan metode yang digunakan\n"
            "□ Data dan sumber yang digunakan\n\n"
            "Juga periksa: apakah metode penilaian sesuai dengan jenis objek dan tujuan penilaian? "
            "Apakah pengungkapan dan asumsi sudah lengkap dan logis?"
        ),
    },
    "🏢 Multi-Objek": {
        "key": "multiobj",
        "desc": "Laporan dengan banyak properti/objek.",
        "instruction": (
            "Ini laporan multi-objek/multi-properti.\n\n"
            "LANGKAH 1 — IDENTIFIKASI: Temukan semua objek properti yang ada. Buat daftar.\n"
            "LANGKAH 2 — PER OBJEK: Untuk SETIAP objek, cek konsistensi secara terpisah "
            "(nama, lokasi, luas, nilai).\n"
            "LANGKAH 3 — ANTAR OBJEK: Apakah ada data objek yang tercampur? "
            "(luas tanah objek A disebut di bagian objek B?)\n"
            "LANGKAH 4 — LINTAS DOKUMEN: Jalankan 3 lapisan pemeriksaan dari Deep Audit "
            "untuk keseluruhan dokumen.\n\n"
            "Tandai setiap temuan dengan nama objek yang relevan di field 'property'."
        ),
    },
    "🏭 Penilaian Aset": {
        "key": "aset",
        "desc": "Laporan penilaian aset/properti (tanah, bangunan, mesin, kendaraan).",
        "instruction": (
            "Ini adalah laporan PENILAIAN ASET/PROPERTI. Lakukan audit komprehensif:\n\n"

            "=== BLOK 1: KONSISTENSI IDENTITAS ===\n"
            "1. Nama obyek & pemilik harus SAMA PERSIS di cover, surat pengantar, "
            "ringkasan eksekutif, uraian umum, dan kesimpulan. "
            "KHUSUS: Cek copy-paste nama dari laporan lain.\n"
            "2. Nomor laporan, tanggal penilaian, tanggal inspeksi, tanggal laporan "
            "harus konsisten dan logis.\n"
            "3. Identitas Pemberi Tugas & Pengguna Laporan harus sama di semua bagian.\n\n"

            "=== BLOK 2: KONSISTENSI LUAS & FISIK ===\n"
            "4. Luas tanah total harus sama di semua bagian. "
            "Jumlah sertifikat × luas per sertifikat = total luas tanah.\n"
            "5. Luas bangunan total & per bangunan/blok harus konsisten.\n"
            "6. Jumlah & nomor sertifikat (SHGB/SHM/SHGU) harus sama.\n"
            "7. Spesifikasi mesin (nama, tipe, tahun, kapasitas): tabel vs uraian.\n"
            "8. Spesifikasi kendaraan (nopol, merk, tipe, tahun): tabel vs uraian.\n\n"

            "=== BLOK 3: KONSISTENSI NILAI ===\n"
            "9. Nilai akhir (Rp & USD) harus SAMA PERSIS di surat pengantar, "
            "resume, ringkasan eksekutif, dan kesimpulan.\n"
            "10. Nilai per komponen (tanah, bangunan, sarana, mesin, kendaraan) "
            "di resume = tabel ringkasan = uraian detail.\n"
            "11. Konversi kurs: Rp ÷ kurs BI = USD. Kurs harus sama di semua bagian.\n"
            "12. BPB × (1 - penyusutan%) = nilai pasar per komponen.\n\n"

            "=== BLOK 4: KESESUAIAN STANDAR ===\n"
            "13. Pendekatan per komponen: tanah→pasar, bangunan/sarana→biaya, "
            "mesin→biaya, kendaraan→pasar. Konsisten antara narasi dan tabel.\n"
            "14. Elemen wajib KEPI & SPI: pernyataan penilai, asumsi, tujuan, "
            "definisi nilai, tanggal penilaian, pendekatan, data yang digunakan, "
            "kejadian penting setelah tanggal penilaian.\n\n"

            "=== BLOK 5: DETEKSI KHUSUS ===\n"
            "15. Artefak copy-paste: nama perusahaan/properti dari laporan lain.\n"
            "16. Placeholder: nilai 'xx', tabel kosong, '...' sebagai pengganti angka.\n"
            "17. Penomoran ganda: tabel atau gambar yang nomornya dipakai dua kali.\n"
            "18. Footer/header salah: menyebut nama laporan/obyek lain.\n"
            "19. Terbilang (huruf) harus sesuai angka di surat pengantar & kesimpulan.\n"
            "20. Format angka konsisten (Rp Ribu vs Rp 000,00 vs US$ ,00)."
        ),
    },
    "📈 Penilaian Saham": {
        "key": "saham",
        "desc": "Laporan penilaian saham (Business Valuation) sesuai POJK 35/2020.",
        "instruction": (
            "Ini adalah laporan PENILAIAN SAHAM (Business Valuation). "
            "Lakukan audit komprehensif dengan SANGAT DETAIL:\n\n"

            "=== BLOK 1: IDENTITAS & KONSISTENSI NAMA ===\n"
            "1. Nama perusahaan yang dinilai dan persentase saham "
            "(misalnya '99,995%') harus KONSISTEN di: cover, surat pengantar, "
            "ringkasan eksekutif, deskripsi penugasan, profil perusahaan, dan kesimpulan. "
            "KHUSUS: Cek apakah ada nama perusahaan/saham LAIN yang tidak relevan "
            "(artefak copy-paste dari laporan lain — ini kesalahan paling umum).\n"
            "2. Nama & singkatan perusahaan yang dinilai: pastikan definisi akronim "
            "muncul saat pertama kali disebut, lalu konsisten digunakan.\n"
            "3. Nama Pemberi Tugas dan Pengguna Laporan: nama, alamat, bidang usaha, "
            "nomor telepon, email, website harus SAMA PERSIS di semua bagian.\n"
            "4. Nomor laporan dan tanggal: nomor laporan, tanggal penerbitan, "
            "dan tanggal penilaian harus sama di semua bagian.\n\n"

            "=== BLOK 2: KONSISTENSI NILAI ===\n"
            "5. Nilai kesimpulan penilaian harus SAMA PERSIS di: surat pengantar, "
            "ringkasan eksekutif (Bab 1), dan kesimpulan (Bab terakhir). "
            "Cek konversi kurs: nilai USD × kurs BI = nilai Rupiah.\n"
            "6. Angka laporan posisi keuangan: Total Aset = Total Liabilitas + Ekuitas "
            "di setiap tahun yang disajikan. Cek juga konsistensi antara tabel dan narasi.\n"
            "7. Aset non-operasional vs operasional: pemisahan konsisten antara "
            "narasi dan tabel. Cek: Nilai 100% ekuitas = Indikasi operasional + Aset non-op.\n"
            "8. Diskon likuiditas pasar: jika diterapkan, pastikan besaran (%) konsisten "
            "antara narasi Bab 1, tabel, dan kesimpulan. Jika tidak diterapkan, "
            "pastikan alasannya konsisten.\n"
            "9. Perhitungan 99,99x%: nilai 99,99x% saham = nilai 100% × 99,99x%. "
            "Cek apakah perhitungan ini konsisten.\n\n"

            "=== BLOK 3: KONSISTENSI METODE & PENDEKATAN ===\n"
            "10. Nama pendekatan & metode harus konsisten di cover letter, Bab 1, "
            "Bab Pendekatan, dan Bab Kesimpulan. "
            "KHUSUS: Cek apakah ada metode/pendekatan untuk perusahaan LAIN "
            "(artefak copy-paste — sangat umum).\n"
            "11. Alasan tidak menggunakan metode lain harus mengacu pada "
            "karakteristik obyek yang benar.\n\n"

            "=== BLOK 4: KESESUAIAN POJK 35/2020 & KEPI/SPI ===\n"
            "12. Elemen wajib POJK 35/2020: (a) status penilai & STTD OJK, "
            "(b) tujuan & maksud penilaian, (c) tanggal efektif penilaian, "
            "(d) dasar nilai, (e) premis penilaian, (f) kondisi pembatas, "
            "(g) asumsi & asumsi khusus, (h) pernyataan independensi, "
            "(i) pernyataan penilai, (j) kejadian setelah tanggal penilaian.\n"
            "13. Premis 'going concern' atau 'likuidasi' harus konsisten "
            "antara pernyataan dan metode yang dipilih.\n\n"

            "=== BLOK 5: TYPO & INKONSISTENSI PENULISAN ===\n"
            "14. TYPO KARAKTER — scan setiap kata dalam dokumen:\n"
            "    - Huruf ganda tidak lazim: 'berrdasarkan', 'deengan', 'adaalah', "
            "'bahwwa', 'terrmasuk', 'untukk', 'peruusahaan', dll.\n"
            "    - Huruf hilang: 'berdsarkan', 'penilaian' → 'penilain', dll.\n"
            "    - Kata duplikat berturut-turut: 'bahwa bahwa', 'yang yang', dll.\n"
            "    Laporkan SETIAP typo yang ditemukan sebagai temuan minor tersendiri "
            "dengan kutipan kalimat dan koreksi yang benar.\n"
            "15. Typo nama perusahaan: ejaan harus konsisten di seluruh dokumen.\n"
            "16. Akronim/singkatan: setiap akronim yang didefinisikan harus "
            "digunakan konsisten setelahnya.\n"
            "17. Paragraf berulang: laporan sering mengulang teks di beberapa bab "
            "(latar belakang, deskripsi penugasan, dll). Pastikan IDENTIK.\n"
            "18. Format angka: konsisten (USD ribu / Rp juta / Rp miliar).\n"
            "19. 'batu bara' vs 'batubara', 'Jl.' vs 'Jalan' — pilih satu format.\n\n"

            "=== BLOK 6: DETEKSI KHUSUS ===\n"
            "19. Artefak copy-paste: cari nama entitas/paragraf dari laporan lain "
            "yang tidak relevan (misalnya 'AMP', 'SBPL', 'PSS', dll).\n"
            "20. Placeholder belum diisi: 'Rp xx juta', tabel dengan hanya kata 'Tabel', "
            "nilai '...' sebagai pengganti angka.\n"
            "21. Penomoran ganda: tabel atau gambar dengan nomor yang sama di dua tempat.\n"
            "22. Footer/header salah: bab dengan footer nama laporan lain.\n"
            "23. Alamat yang disebut dua kali berturut-turut tanpa pemisah yang jelas.\n\n"

            "Berikan temuan SPESIFIK dengan menyebutkan halaman/bab dan "
            "nilai/kata yang tidak konsisten. Prioritaskan temuan yang mempengaruhi "
            "validitas penilaian."
        ),
    },
    "⚖️ Pendapat Kewajaran": {
        "key": "fairness",
        "desc": "Laporan pendapat kewajaran (Fairness Opinion) atas transaksi.",
        "instruction": (
            "Ini adalah laporan PENDAPAT KEWAJARAN (Fairness Opinion). "
            "Lakukan audit komprehensif:\n\n"

            "=== BLOK 1: KONSISTENSI IDENTITAS TRANSAKSI ===\n"
            "1. Nama rencana transaksi harus KONSISTEN di seluruh laporan.\n"
            "2. Nama penjual, pembeli, dan obyek transaksi harus SAMA PERSIS. "
            "KHUSUS: Cek copy-paste nama dari laporan lain.\n"
            "3. Nilai transaksi (harga per saham, total) harus sama di semua bagian "
            "dan konsisten dengan hasil penilaian.\n\n"

            "=== BLOK 2: KONSISTENSI ANALISIS ===\n"
            "4. Kewajaran finansial: analisis konsisten antara narasi dan kesimpulan.\n"
            "5. Kewajaran non-finansial: konsisten antara analisis dan kesimpulan.\n"
            "6. Referensi ke laporan penilaian: nilai yang dikutip harus sama.\n\n"

            "=== BLOK 3: KESESUAIAN STANDAR ===\n"
            "7. Elemen wajib Pendapat Kewajaran POJK 35/2020: identitas peminta, "
            "deskripsi transaksi, analisis kewajaran finansial & non-finansial, "
            "kesimpulan, pernyataan independensi, kondisi pembatas.\n"
            "8. Pengungkapan afiliasi jika transaksi dengan pihak terafiliasi.\n\n"

            "=== BLOK 4: DETEKSI KHUSUS ===\n"
            "9. Artefak copy-paste, placeholder, penomoran ganda, footer salah.\n"
            "10. Konsistensi paragraf berulang, format angka, akronim."
        ),
    },
}

# Check items per mode
CHECK_ITEMS_DEFAULT = [
    "Konsistensi Tanggal (inspeksi, penilaian, laporan)",
    "Konsistensi Luas (tanah, bangunan, GFA, NLA)",
    "Konsistensi Alamat & Lokasi",
    "Konsistensi Nilai (angka vs huruf, ringkasan vs kesimpulan)",
    "Kepemilikan & Nomor Sertifikat",
    "Kesesuaian Standar KEPI/MAPPI",
    "Analisis Pasar & Data Pembanding",
    "Pendekatan & Metode Penilaian",
    "Kelengkapan Narasi & Deskripsi Objek",
    "Deteksi Artefak Copy-Paste",
    "Deteksi Placeholder Belum Diisi",
    "Konsistensi Penomoran Tabel/Gambar/Sub-bab",
]

CHECK_ITEMS_SAHAM = [
    "Konsistensi nama & persentase saham di seluruh laporan",
    "Deteksi nama entitas/artefak copy-paste dari laporan lain",
    "Nilai kesimpulan konsisten di surat pengantar, ringkasan, dan kesimpulan",
    "Konversi kurs BI (nilai USD × kurs = nilai Rupiah)",
    "Angka LK: Total Aset = Total Liabilitas + Ekuitas",
    "Konsistensi aset operasional vs non-operasional",
    "Konsistensi diskon likuiditas pasar",
    "Konsistensi Pemberi Tugas & Pengguna Laporan",
    "Konsistensi nomor laporan & tanggal penilaian",
    "Konsistensi pendekatan & metode di semua bab",
    "Kelengkapan elemen wajib POJK 35/2020",
    "Konsistensi akronim & singkatan",
    "Konsistensi paragraf berulang antar bab",
    "Format angka konsisten (USD ribu / Rp juta)",
    "Deteksi placeholder belum diisi (Rp xx, tabel kosong)",
    "Deteksi penomoran ganda tabel/gambar",
    "Deteksi footer/header mengacu laporan lain",
    "Konsistensi ejaan 'batu bara' vs 'batubara'",
]

CHECK_ITEMS_ASET = [
    "Konsistensi nama obyek penilaian & pemilik",
    "Deteksi copy-paste nama dari laporan lain",
    "Konsistensi nomor laporan, tanggal penilaian, inspeksi, dan terbit",
    "Konsistensi Pemberi Tugas & Pengguna Laporan",
    "Konsistensi luas tanah total & per sertifikat",
    "Konsistensi luas bangunan total & per bangunan",
    "Konsistensi nomor, tanggal terbit/berakhir sertifikat",
    "Nilai kesimpulan (Rp & USD) konsisten di semua bagian",
    "Nilai per komponen konsisten (tanah, bangunan, sarana, mesin, kendaraan)",
    "Konversi kurs BI (Rp ÷ kurs = USD)",
    "Verifikasi BPB × (1 - penyusutan%) = nilai pasar",
    "Pemisahan mesin DIGUNAKAN vs BELUM DIGUNAKAN",
    "Spesifikasi mesin (nama, tipe, tahun, kapasitas) tabel vs uraian",
    "Spesifikasi kendaraan (nopol, merk, tipe, tahun) tabel vs uraian",
    "Pendekatan penilaian per komponen (pasar/biaya/pendapatan)",
    "Terbilang (huruf) sesuai angka",
    "Kelengkapan elemen wajib KEPI & SPI",
    "Deteksi placeholder belum diisi",
    "Deteksi penomoran ganda tabel/gambar/sub-bab",
    "Deteksi footer/header mengacu laporan lain",
]

CHECK_ITEMS_FAIRNESS = [
    "Konsistensi nama & deskripsi rencana transaksi",
    "Konsistensi nama pihak-pihak transaksi",
    "Deteksi copy-paste dari laporan lain",
    "Konsistensi nilai transaksi",
    "Referensi silang dengan laporan penilaian yang dirujuk",
    "Konsistensi kesimpulan kewajaran",
    "Kelengkapan elemen wajib POJK 35/2020",
    "Pengungkapan hubungan afiliasi",
    "Konsistensi tanggal",
    "Deteksi placeholder & penomoran ganda",
]

# ──────────────────────────────────────────────
# SYSTEM PROMPT
# ──────────────────────────────────────────────
SYSTEM_PROMPT = """Kamu adalah EDITOR SENIOR laporan penilaian di KJPP Suwendho Rinaldy dan Rekan (KJPP SRR).
Kamu bukan mesin pencari kata kunci. Kamu membaca laporan dengan pikiran kritis penuh, persis seperti
editor berpengalaman yang mempertanyakan setiap kata, angka, dan paragraf yang kamu temui.

═══════════════════════════════════════════════════════
MENTALITAS EDITOR: BACA INI SEBELUM APAPUN
═══════════════════════════════════════════════════════

Setiap kali kamu menyentuh sebuah kata, kalimat, angka, atau paragraf — tanya pada dirimu:
"Apakah ini terasa benar? Apakah ini masuk akal? Apakah ini konsisten dengan yang sudah kubaca?"

PRINSIP UTAMA:
1. Typo di setiap laporan adalah UNIK — jangan hanya mengandalkan daftar pola.
   Gunakan intuisi bahasamu untuk mendeteksi apapun yang terasa "ganjil".
2. Jika terasa aneh, CATAT. Lebih baik 5 temuan yang ternyata bukan masalah
   daripada melewatkan 1 kesalahan fatal.
3. Verifikasi silang aktif — saat membaca Bab 4, ingat apa yang ada di Bab 1.
4. Baca HEADER dan FOOTER setiap halaman — ini sering mengandung kesalahan serius.
5. Bangun model mental: siapa nama perusahaan/objek, bisnisnya apa, di mana,
   siapa kliennya, berapa nilainya. Setiap informasi baru — cek konsistensinya.

═══════════════════════════════════════════════════════
TIGA LAPISAN PEMERIKSAAN (JALANKAN SIMULTAN)
═══════════════════════════════════════════════════════

LAPISAN 1 — LINGUISTIK (kata demi kata):
• Apakah setiap kata terasa utuh dan benar secara ejaan?
• Kata terulang dua kali berturutan? (mis: "terdiri terdiri dari")
• Huruf visual mirip tapi beda: l kecil vs I kapital, 0 vs O, rn vs m, cl vs d
• Tanda baca menggantikan huruf: "era!" → seharusnya "erat", "1nflasi" → "Inflasi"
• Spasi tersesat di tengah kata: "tu run" → "turun", "ber sama" → "bersama"
• Awal kalimat tidak kapital, kalimat tidak berakhir tanda titik
• Format angka Indonesia: titik=ribuan (Rp1.250.000), koma=desimal (10,5%)
  SALAH: "Rp1,250,000" atau "Rp. 1.250.000" atau "Rp 283.118, juta"
• Huruf ganda tidak lazim: "berrdasarkan", "secarra", "prusahaan", "penigkatan"
• Ejaan KBBI: "batu bara" (dua kata), "bertanggung jawab" (dua kata),
  "kerja sama" (dua kata), "izin" bukan "ijin", "praktik" bukan "praktek"

LAPISAN 2 — KONTEKS & RELEVANSI (paragraf demi paragraf):
• Apakah semua nama entitas/singkatan sudah diperkenalkan dalam laporan ini?
  Nama muncul tiba-tiba tanpa penjelasan = kemungkinan besar sisa laporan lain → 🔴 KRITIS
• Konteks narasi relevan dengan bisnis/aset yang dinilai?
  (laporan pertambangan batu bara di Kalimantan tidak membahas "pasar Singapura")
• Footer/header setiap halaman — apakah menyebut nama proyek yang benar?
  Footer salah muncul di 6 halaman berturut-turut = 🔴 KRITIS
• Placeholder belum diisi: "Rp xx", "[nilai]", "[nama]", tabel berisi kata "Tabel" saja = 🔴 KRITIS

LAPISAN 3 — KONSISTENSI (dokumen secara keseluruhan):
• Angka yang sama (nilai, luas, %) — selalu sama di setiap bagian?
• Angka dalam narasi = angka dalam tabel?
• Nilai di kesimpulan = nilai di bab penilaian = nilai di ringkasan eksekutif?
• Nama perusahaan/objek — ejaan sama persis di seluruh dokumen?
• Nomor sub-bab urut dan logis? (3.6.2 di Bab 3.7 = kesalahan)
• Nomor gambar/tabel urut, tidak ada duplikat?
• Judul sub-bab mencerminkan isinya?
• Cover dan footer menyebut jenis laporan yang sama?
• Urutan kronologis masuk akal: tanggal inspeksi ≤ tanggal cut-off ≤ tanggal laporan?

═══════════════════════════════════════════════════════
STRUKTUR OUTPUT — BAGIAN A SAMPAI J
═══════════════════════════════════════════════════════

Kelompokkan temuan ke dalam bagian berikut menggunakan field "category":
• "Konsistensi Nilai" / "Konsistensi Tanggal" / "Konsistensi Luas" → Bagian A
• "Artefak Copy-Paste" → Bagian B
• "Placeholder" → Bagian C
• "Penomoran" → Bagian D
• "Nama Perusahaan" → Bagian E
• "Ejaan & Penulisan" → Bagian F
• "Typo" → Bagian G
• "Logis/Substansi" → Bagian H
• "Daftar Isi" / "Daftar Gambar" / "Daftar Tabel" → Bagian I
• "KEPI/SPI/POJK" → Bagian J

═══════════════════════════════════════════════════════
PANDUAN SEVERITY
═══════════════════════════════════════════════════════

🔴 kritis  — mempengaruhi validitas atau dapat menyesatkan pembaca secara signifikan:
             artefak copy-paste, placeholder, footer salah, angka material tidak konsisten
🟠 mayor   — masalah serius yang harus diperbaiki sebelum terbit:
             inkonsistensi angka non-material, nama entitas tidak konsisten,
             judul sub-bab tidak sesuai isi, penomoran salah
🟡 minor   — ketidakkonsistenan kecil, typo, inkonsistensi ejaan
🟢 ok      — elemen yang sudah benar dan sesuai standar (gunakan jika perlu dikonfirmasi)
🔵 info    — catatan atau saran perbaikan

═══════════════════════════════════════════════════════
FORMAT JSON OUTPUT (WAJIB — TIDAK BOLEH ADA TEKS DI LUAR JSON)
═══════════════════════════════════════════════════════

{
  "report_type": "deskripsi jenis laporan (mis: Laporan Penilaian Saham PT KMS)",
  "valuation_type": "Saham|Aset|Perkebunan|Hotel|Tanah & Bangunan|Fairness Opinion",
  "properties": ["nama entitas/objek utama yang dinilai"],
  "summary": {
    "total_findings": 0,
    "kritis": 0,
    "mayor": 0,
    "minor": 0,
    "ok": 0,
    "overall_score": 85,
    "status": "LULUS|PERLU REVISI MINOR|PERLU REVISI MAYOR",
    "executive_summary": "2-3 kalimat ringkasan kritis: apa masalah terbesar, dampaknya apa, rekomendasi utama."
  },
  "findings": [
    {
      "id": "F001",
      "severity": "kritis",
      "category": "Artefak Copy-Paste",
      "title": "Judul singkat (maks 10 kata)",
      "detail": "Penjelasan detail dengan kutipan teks bermasalah dan lokasi spesifik.",
      "teks_salah": "kutipan teks yang salah",
      "seharusnya": "saran perbaikan / teks yang benar",
      "page_hint": "Hal. 14 / Bab 1.7",
      "property": ""
    }
  ]
}

CATATAN PENTING:
- Untuk kategori "Typo": WAJIB isi teks_salah dan seharusnya
- Untuk kategori "Konsistensi Nilai": isi teks_salah dengan nilai yang salah, seharusnya dengan nilai yang benar
- Sebutkan nomor halaman dan/atau bab yang spesifik di page_hint
- Kutip teks asli dalam tanda petik di field detail
- overall_score: 100 = sempurna, kurangi ~10 per temuan kritis, ~5 per mayor, ~2 per minor"""


# ══════════════════════════════════════════════
# EXCEL PARSER
# ══════════════════════════════════════════════

def parse_excel_workbook(file) -> dict:
    """
    Parse file XLSX dan ekstrak semua data numerik beserta label/konteksnya.
    Returns dict: {
        "sheets": [{name, rows: [{label, value, cell, row_idx, col_idx}]}],
        "number_map": {value_int: [{label, sheet, cell}]},
        "all_numbers": [float],
        "summary": str
    }
    """
    if not EXCEL_AVAILABLE:
        return {"error": "openpyxl tidak tersedia", "sheets": [], "number_map": {}, "all_numbers": []}

    try:
        wb = openpyxl.load_workbook(file, data_only=True)
        result = {"sheets": [], "number_map": {}, "all_numbers": [], "summary": ""}
        summary_lines = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_data = {"name": sheet_name, "rows": []}
            sheet_numbers = 0

            # Baca semua baris
            for row in ws.iter_rows():
                row_label = ""
                # Ambil label dari kolom teks pertama di baris ini
                for cell in row:
                    if cell.value is not None and isinstance(cell.value, str) and cell.value.strip():
                        row_label = cell.value.strip()[:80]
                        break

                for cell in row:
                    if cell.value is None:
                        continue
                    val = cell.value

                    # Proses nilai numerik
                    if isinstance(val, (int, float)) and not isinstance(val, bool):
                        if val == 0:
                            continue
                        # Cari label dari kolom sebelumnya di baris yang sama
                        label = row_label

                        entry = {
                            "label": label,
                            "value": float(val),
                            "cell": cell.coordinate,
                            "sheet": sheet_name,
                        }
                        sheet_data["rows"].append(entry)
                        result["all_numbers"].append(float(val))
                        sheet_numbers += 1

                        # Tambah ke number_map dengan key = round ke integer
                        key = int(abs(val))
                        if key > 0:
                            if key not in result["number_map"]:
                                result["number_map"][key] = []
                            result["number_map"][key].append({
                                "label": label,
                                "sheet": sheet_name,
                                "cell": cell.coordinate,
                                "value": float(val),
                            })

            result["sheets"].append(sheet_data)
            if sheet_numbers > 0:
                summary_lines.append(f"Sheet '{sheet_name}': {sheet_numbers} nilai numerik")

        result["summary"] = "\n".join(summary_lines) if summary_lines else "Tidak ada data numerik ditemukan"
        return result

    except Exception as e:
        return {"error": str(e), "sheets": [], "number_map": {}, "all_numbers": []}


def extract_doc_numbers(doc_text: str) -> list:
    """
    Ekstrak semua angka dari teks dokumen beserta konteksnya.
    Returns list of {value, context, raw_str}
    """
    results = []
    # Pattern: angka dengan titik sebagai ribuan dan koma sebagai desimal (format Indonesia)
    pattern = re.compile(
        r'(?<![.\d])'                     # tidak diawali digit/titik
        r'(\d{1,3}(?:\.\d{3})+(?:,\d+)?'  # 1.234.567 atau 1.234.567,89
        r'|\d{4,}(?:,\d+)?)'              # atau 4+ digit tanpa titik
        r'(?![.\d])',                      # tidak diikuti digit/titik
        re.MULTILINE
    )

    for m in pattern.finditer(doc_text):
        raw = m.group()
        # Parse ke float
        try:
            # Format Indonesia: titik = ribuan, koma = desimal
            normalized = raw.replace('.', '').replace(',', '.')
            val = float(normalized)
            if val < 10:   # skip angka kecil
                continue
            # Ambil konteks 60 karakter di sekitar angka
            start = max(0, m.start() - 60)
            end   = min(len(doc_text), m.end() + 60)
            ctx   = doc_text[start:end].replace('\n', ' ').strip()
            results.append({"value": val, "raw_str": raw, "context": ctx})
        except (ValueError, OverflowError):
            continue

    # Deduplikasi berdasarkan value + context yang mirip
    seen = set()
    deduped = []
    for r in results:
        key = (int(r["value"]), r["context"][:30])
        if key not in seen:
            seen.add(key)
            deduped.append(r)
    return deduped


def compare_doc_with_excel(doc_numbers: list, excel_data: dict,
                           tolerance_pct: float = 0.001) -> dict:
    """
    Bandingkan angka di dokumen dengan data Excel.
    Returns {
        "matches": [...],
        "not_found": [...],
        "summary": str
    }
    """
    if not excel_data or not excel_data.get("number_map"):
        return {"matches": [], "not_found": [], "summary": "Data Excel tidak tersedia"}

    number_map = excel_data["number_map"]
    matches    = []
    not_found  = []

    for dn in doc_numbers:
        val      = dn["value"]
        val_int  = int(abs(val))
        found    = False

        # Cari exact match (dalam toleransi)
        for delta in range(-2, 3):  # toleransi ±2 karena pembulatan
            key = val_int + delta
            if key in number_map:
                for em in number_map[key]:
                    diff_pct = abs(val - em["value"]) / max(abs(em["value"]), 1)
                    if diff_pct <= tolerance_pct:
                        matches.append({
                            "doc_value":   val,
                            "doc_raw":     dn["raw_str"],
                            "doc_context": dn["context"],
                            "excel_value": em["value"],
                            "excel_label": em["label"],
                            "excel_sheet": em["sheet"],
                            "excel_cell":  em["cell"],
                            "status":      "match",
                        })
                        found = True
                        break
                if found:
                    break

        if not found and val > 100:  # hanya flag angka > 100
            not_found.append({
                "doc_value":   val,
                "doc_raw":     dn["raw_str"],
                "doc_context": dn["context"],
                "status":      "not_found",
            })

    n_match = len(matches)
    n_miss  = len(not_found)
    total   = n_match + n_miss
    pct     = (n_match / total * 100) if total > 0 else 0

    summary = (
        f"Total angka diperiksa: {total} | "
        f"✅ Cocok: {n_match} ({pct:.0f}%) | "
        f"❓ Tidak ditemukan di Excel: {n_miss}"
    )
    return {"matches": matches, "not_found": not_found, "summary": summary}


def build_excel_context_for_prompt(excel_data: dict, max_chars: int = 8000) -> str:
    """
    Buat ringkasan data Excel untuk disertakan dalam prompt Claude.
    """
    if not excel_data or excel_data.get("error"):
        return ""

    lines = ["=== DATA LEMBAR KERJA (EXCEL) ===\n"]
    lines.append(excel_data.get("summary", "") + "\n")

    total_chars = sum(len(l) for l in lines)
    for sheet in excel_data.get("sheets", []):
        sname = sheet["name"]
        rows  = sheet["rows"]
        if not rows:
            continue
        lines.append(f"\n--- Sheet: {sname} ---\n")
        for r in rows:
            if r["value"] == 0:
                continue
            line = f"  [{r['cell']}] {r['label']}: {r['value']:,.2f}\n"
            if total_chars + len(line) > max_chars:
                lines.append("  ... (data terpotong)\n")
                break
            lines.append(line)
            total_chars += len(line)

    lines.append("\nInstruksi komparasi Excel:\n")
    lines.append(
        "Bandingkan angka-angka penting dalam laporan dengan data Excel di atas. "
        "Jika ada angka di laporan yang tidak cocok dengan Excel, atau ada angka di Excel "
        "yang tidak ada di laporan, tandai sebagai temuan dengan severity 'kritikal'. "
        "Khususnya periksa: nilai kesimpulan penilaian, angka laporan keuangan, "
        "persentase kepemilikan, dan nilai per komponen/aset.\n"
    )
    return "".join(lines)


# ══════════════════════════════════════════════
# LOKAL CHECKER: ARTEFAK & PLACEHOLDER
# ══════════════════════════════════════════════

def cek_typo_lokal(doc_text: str) -> list:
    """
    Deteksi typo lokal menggunakan regex tanpa AI.
    Setiap pola menghasilkan SATU finding per kemunculan unik,
    dengan field teks_salah dan seharusnya untuk tampilan tabel.
    """
    findings = []
    lt_counter = [0]  # mutable counter untuk id unik

    def make_typo_finding(teks_salah: str, seharusnya: str, keterangan: str,
                          page_hint: str = "Seluruh dokumen") -> dict:
        lt_counter[0] += 1
        return {
            "id": f"LT{lt_counter[0]:02d}",
            "severity": "minor",
            "category": "Typo",
            "title": teks_salah,
            "teks_salah": teks_salah,
            "seharusnya": seharusnya,
            "detail": keterangan,
            "page_hint": page_hint,
            "property": "",
        }

    # ── POLA TYPO SPESIFIK ────────────────────────────────────────────────────
    # Format: (regex_pola, teks_salah_label, seharusnya, keterangan)
    SPECIFIC_TYPOS = [
        # Huruf ganda
        (r'\bberrdasarkan\b',          '"berrdasarkan"',        '"berdasarkan"',                 'huruf "r" ganda'),
        (r'\bsecarra\b',               '"secarra"',             '"secara"',                      'huruf "r" ganda'),
        (r'\bperuusahaan\b',           '"peruusahaan"',         '"perusahaan"',                  'huruf "u" ganda'),
        (r'\bperrseroan\b',            '"perrseroan"',          '"perseroan"',                   'huruf "r" ganda'),
        (r'\bterrsebut\b',             '"terrsebut"',           '"tersebut"',                    'huruf "r" ganda'),
        (r'\bterrmasuk\b',             '"termasuk"',            '"termasuk"',                    'huruf "r" ganda'),
        (r'\bdeengan\b',               '"deengan"',             '"dengan"',                      'huruf "e" ganda'),
        (r'\bdaalam\b',                '"daalam"',              '"dalam"',                       'huruf "a" ganda'),
        (r'\badaalah\b',               '"adaalah"',             '"adalah"',                      'huruf "a" ganda'),
        (r'\bmenngenai\b',             '"menngenai"',           '"mengenai"',                    'huruf "n" ganda'),
        (r'\bperrtimbangan\b',         '"perrtimbangan"',       '"pertimbangan"',                'huruf "r" ganda'),
        (r'\bnilaai\b',                '"nilaai"',              '"nilai"',                       'huruf "a" ganda'),
        # Huruf hilang
        (r'\bprusahaan\b',             '"prusahaan"',           '"perusahaan"',                  'huruf "e" hilang'),
        (r'\bpenigkatan\b',            '"penigkatan"',          '"peningkatan"',                 'huruf "n" hilang'),
        (r'\bmenujukkan\b',            '"menujukkan"',          '"menunjukkan"',                 'huruf "n" hilang'),
        (r'\bdiperlkukan\b',           '"diperlkukan"',         '"diperlukan"',                  'huruf "u" hilang'),
        (r'\bberdsarkan\b',            '"berdsarkan"',          '"berdasarkan"',                 'huruf "a" hilang'),
        (r'\bpenilain\b',              '"penilain"',            '"penilaian"',                   'huruf "ia" hilang'),
        (r'\bkeuangn\b',               '"keuangn"',             '"keuangan"',                    'huruf "a" hilang'),
        (r'\blnflasi\b',               '"lnflasi"',             '"Inflasi"',                     'huruf "I" kapital ditulis "l" kecil'),
        (r'\b\(ima\s+ratus\b',         '"(ima ratus"',          '"(lima ratus"',                 'huruf "l" hilang'),
        # Kata duplikat berturut-turut
        (r'\bterdiri\s+terdiri\b',     '"terdiri terdiri"',     '"terdiri"',                     'kata duplikat berturut-turut'),
        (r'\bperbandingan\s+perbandingan\b', '"perbandingan perbandingan"', '"perbandingan"',    'kata duplikat berturut-turut'),
        (r'\b(\w{4,})\s+\1\b',         '[kata berulang]',       '[hilangkan duplikat]',          'kata yang sama berulang berturut-turut'),
        # Karakter salah
        (r'\bsinergi\s+era!',          '"sinergi era!"',        '"sinergi erat"',                'tanda seru menggantikan huruf "t"'),
        (r'\btu\s+run\b',              '"tu run"',              '"turun"',                       'spasi di tengah kata "turun"'),
        # Format angka / tanda baca salah
        (r'Rp\s*[\d\.]+,\s*juta',      '"Rp xxx, juta"',        '"Rp xxx juta"',                 'koma di posisi salah antara angka dan "juta"'),
        # Awal kalimat tidak kapital setelah titik (sampling pada "pada 20xx")
        (r'(?<=[.!?] )pada\s+20\d\d',         '"pada 20xx" (awal kalimat)', '"Pada 20xx"',       'awal kalimat harus kapital'),
        # Spasi ganda antar kata
        (r'\bPT\s{2,}(?=[A-Z])',       '"PT  [nama]"',          '"PT [nama]"',                   'spasi ganda setelah "PT"'),
        # Kata umum lainnya
        (r'\bmenuruurut\b',            '"menuruurut"',          '"menurut"',                     'huruf "u" berlebih'),
        (r'\btahunm\b',                '"tahunm"',              '"tahun"',                       'huruf "m" berlebih'),
        (r'\bpemilkk?\b',              '"pemilk"',              '"pemilik"',                     'huruf "ik" salah'),
        (r'\bpenilaiiaan\b',           '"penilaiiaan"',         '"penilaian"',                   'huruf "i" & "a" ganda'),
    ]

    seen_patterns: set = set()  # hindari duplikat pola yang sama

    for pat, teks_salah_label, seharusnya_label, keterangan in SPECIFIC_TYPOS:
        for m in re.finditer(pat, doc_text, re.IGNORECASE):
            matched = m.group()
            # Untuk pola generik (kata duplikat), gunakan teks asli
            if teks_salah_label == '[kata berulang]':
                teks_salah_label = f'"{matched}"'
                seharusnya_label = f'"{m.group(1)}" (hapus duplikat)'

            dedup_key = pat + "::" + matched.lower()
            if dedup_key in seen_patterns:
                continue
            seen_patterns.add(dedup_key)

            # Ambil konteks (~80 karakter di kiri/kanan)
            ctx_start = max(0, m.start() - 80)
            ctx_end   = min(len(doc_text), m.end() + 80)
            ctx = doc_text[ctx_start:ctx_end].replace('\n', ' ').strip()

            findings.append(make_typo_finding(
                teks_salah  = teks_salah_label,
                seharusnya  = seharusnya_label,
                keterangan  = f'{keterangan} — konteks: "…{ctx}…"',
                page_hint   = "Cek seluruh dokumen",
            ))

    return findings


def cek_artefak_dan_placeholder(doc_text: str, mode_key: str) -> dict:
    """
    Deteksi lokal (tanpa AI) untuk:
    1. Artefak copy-paste (nama entitas asing)
    2. Placeholder belum diisi
    3. Penomoran ganda tabel/gambar
    4. Typo (huruf ganda, salah ketik, spasi ganda)
    """
    findings = []

    # ── CEK PLACEHOLDER ──────────────────────────────────────────────────────
    ph_hits = []
    for pat_str in PLACEHOLDER_PATTERNS:
        for m in re.finditer(pat_str, doc_text, re.IGNORECASE | re.MULTILINE):
            ctx_start = max(0, m.start() - 80)
            ctx_end   = min(len(doc_text), m.end() + 80)
            ctx = doc_text[ctx_start:ctx_end].replace('\n', ' ').strip()
            ph_hits.append({"text": m.group(), "context": ctx})

    if ph_hits:
        hits_str = "\n".join(
            f'• "{h["text"]}" → ...{h["context"]}...'
            for h in ph_hits[:8]
        )
        findings.append({
            "id": "L001",
            "severity": "kritikal",
            "category": "Placeholder",
            "title": f"Ditemukan {len(ph_hits)} placeholder yang belum diisi",
            "detail": (
                f"Terdapat {len(ph_hits)} bagian yang tampak belum selesai diisi:\n"
                + hits_str
            ),
            "page_hint": "Cek seluruh dokumen",
            "property": "",
        })
    else:
        findings.append({
            "id": "L001",
            "severity": "ok",
            "category": "Placeholder",
            "title": "Tidak ditemukan placeholder yang belum diisi",
            "detail": "Semua nilai dalam dokumen tampak sudah diisi (tidak ada 'xx', tabel kosong, dll).",
            "page_hint": "",
            "property": "",
        })

    # ── CEK PENOMORAN GANDA TABEL ─────────────────────────────────────────────
    tabel_nums = {}
    for m in re.finditer(r'\bTabel\s+(\d+[\.\d]*)', doc_text, re.IGNORECASE):
        num = m.group(1)
        tabel_nums.setdefault(num, []).append(m.start())
    duplikat_tabel = {k: v for k, v in tabel_nums.items() if len(v) > 2}
    # > 2 karena penomoran mungkin muncul dalam daftar tabel dan di body

    gambar_nums = {}
    for m in re.finditer(r'\bGambar\s+(\d+[\.\d]*)', doc_text, re.IGNORECASE):
        num = m.group(1)
        gambar_nums.setdefault(num, []).append(m.start())
    duplikat_gambar = {k: v for k, v in gambar_nums.items() if len(v) > 2}

    if duplikat_tabel or duplikat_gambar:
        detail_parts = []
        for num, positions in sorted(duplikat_tabel.items()):
            detail_parts.append(f"Tabel {num}: muncul {len(positions)} kali")
        for num, positions in sorted(duplikat_gambar.items()):
            detail_parts.append(f"Gambar {num}: muncul {len(positions)} kali")
        findings.append({
            "id": "L002",
            "severity": "minor",
            "category": "Penomoran",
            "title": f"Penomoran duplikat: {len(duplikat_tabel)} tabel, {len(duplikat_gambar)} gambar",
            "detail": (
                "Nomor tabel/gambar berikut muncul lebih dari 2 kali — "
                "kemungkinan ada penomoran ganda:\n"
                + "\n".join(f"• {d}" for d in detail_parts[:10])
            ),
            "page_hint": "Cek daftar tabel/gambar",
            "property": "",
        })
    else:
        findings.append({
            "id": "L002",
            "severity": "ok",
            "category": "Penomoran",
            "title": "Tidak ditemukan penomoran ganda tabel/gambar",
            "detail": "Penomoran tabel dan gambar tampak unik.",
            "page_hint": "",
            "property": "",
        })

    # ── CEK PENOMORAN SUB-BAB LOMPAT ─────────────────────────────────────────
    subbab_issues = []
    # Cari pola x.y.z dan cek apakah ada lompatan
    section_nums = []
    for m in re.finditer(r'\b(\d+)\.(\d+)\.(\d+)\b', doc_text):
        section_nums.append((int(m.group(1)), int(m.group(2)), int(m.group(3)), m.start()))

    # Deteksi sub-bab yang salah parent (misal 3.6.2 di dalam 3.7)
    for i in range(1, len(section_nums)):
        prev = section_nums[i-1]
        curr = section_nums[i]
        if curr[0] == prev[0] and curr[1] > prev[1] + 1:
            # Ada lompatan di level 2 dengan parent yang sama
            subbab_issues.append(
                f"{prev[0]}.{prev[1]}.x → {curr[0]}.{curr[1]}.{curr[2]} "
                f"(lompat {curr[1] - prev[1] - 1} level)"
            )

    if subbab_issues:
        findings.append({
            "id": "L003",
            "severity": "minor",
            "category": "Penomoran",
            "title": f"Ditemukan {len(subbab_issues)} penomoran sub-bab yang tidak urut",
            "detail": (
                "Penomoran sub-bab berikut tampak tidak urut atau salah:\n"
                + "\n".join(f"• {s}" for s in subbab_issues[:5])
                + "\nPeriksa apakah ada nomor sub-bab yang salah (misalnya 3.6.2 seharusnya 3.7.1)."
            ),
            "page_hint": "Cek daftar isi",
            "property": "",
        })

    # ── CEK KONSISTENSI EJAAN "batu bara" / "batubara" ──────────────────────
    n_batu_bara = len(re.findall(r'\bbatu bara\b', doc_text, re.IGNORECASE))
    n_batubara  = len(re.findall(r'\bbatubara\b',  doc_text, re.IGNORECASE))
    if n_batu_bara > 0 and n_batubara > 0:
        findings.append({
            "id": "L004",
            "severity": "minor",
            "category": "Ejaan & Penulisan",
            "title": "Penulisan 'batu bara' dan 'batubara' tidak konsisten",
            "detail": (
                f"'batu bara' (dua kata) muncul {n_batu_bara}× dan "
                f"'batubara' (satu kata) muncul {n_batubara}×. "
                "Pilih satu format dan terapkan konsisten di seluruh dokumen. "
                "EYD menggunakan 'batu bara' (dua kata)."
            ),
            "page_hint": "Seluruh dokumen",
            "property": "",
        })

    # ── CEK TYPO LOKAL ────────────────────────────────────────────────────────
    typo_findings = cek_typo_lokal(doc_text)
    findings.extend(typo_findings)

    return {
        "findings":   findings,
        "n_kritikal": sum(1 for f in findings if f["severity"] == "kritikal"),
        "n_minor":    sum(1 for f in findings if f["severity"] == "minor"),
        "n_ok":       sum(1 for f in findings if f["severity"] == "ok"),
    }


# ══════════════════════════════════════════════
# GOOGLE SHEETS — CONNECTION & HELPERS
# ══════════════════════════════════════════════

@st.cache_resource(show_spinner=False)
def get_gsheet_client():
    if not GSPREAD_AVAILABLE:
        return None, None, "gspread tidak tersedia"
    if "gcp_service_account" not in st.secrets:
        return None, None, "KEY_MISSING: 'gcp_service_account' tidak ditemukan di Secrets"
    if "spreadsheet_id" not in st.secrets:
        return None, None, "KEY_MISSING: 'spreadsheet_id' tidak ditemukan di Secrets"
    try:
        creds_dict   = dict(st.secrets["gcp_service_account"])
        creds        = Credentials.from_service_account_info(creds_dict, scopes=GSPREAD_SCOPES)
        client       = gspread.authorize(creds)
        spreadsheet  = client.open_by_key(st.secrets["spreadsheet_id"])
        return client, spreadsheet, None
    except Exception as e:
        return None, None, str(e)


def get_or_create_sheet(spreadsheet, sheet_name: str, headers: list):
    try:
        ws = spreadsheet.worksheet(sheet_name)
    except Exception:
        ws = spreadsheet.add_worksheet(title=sheet_name, rows=1000, cols=len(headers))
        ws.append_row(headers)
    existing = ws.row_values(1)
    if existing != headers:
        ws.insert_row(headers, 1)
    return ws


RIWAYAT_HEADERS = [
    "audit_id", "timestamp", "files", "mode",
    "score", "kritikal", "minor",
    "report_type", "properties", "executive_summary", "findings_json"
]


def load_riwayat(spreadsheet) -> dict:
    if spreadsheet is None:
        return st.session_state.get("riwayat_local", {})
    try:
        ws      = get_or_create_sheet(spreadsheet, SHEET_RIWAYAT, RIWAYAT_HEADERS)
        records = ws.get_all_records()
        result  = {}
        for row in records:
            aid = str(row.get("audit_id", ""))
            if not aid:
                continue
            try:
                findings = json.loads(row.get("findings_json", "[]"))
            except Exception:
                findings = []
            try:
                props = json.loads(row.get("properties", "[]"))
            except Exception:
                props = []
            result[aid] = {
                "timestamp": row.get("timestamp", ""),
                "files":     row.get("files", "").split("|"),
                "mode":      row.get("mode", ""),
                "score":     int(row.get("score", 0)),
                "kritikal":  int(row.get("kritikal", 0)),
                "minor":     int(row.get("minor", 0)),
                "result": {
                    "report_type": row.get("report_type", ""),
                    "properties":  props,
                    "summary": {
                        "overall_score":     int(row.get("score", 0)),
                        "kritikal":          int(row.get("kritikal", 0)),
                        "minor":             int(row.get("minor", 0)),
                        "executive_summary": row.get("executive_summary", ""),
                    },
                    "findings": findings,
                },
            }
        return result
    except Exception as e:
        st.warning(f"⚠️ Gagal membaca riwayat: {e}")
        return st.session_state.get("riwayat_local", {})


def save_riwayat_row(spreadsheet, audit_id: str, data: dict):
    if spreadsheet is None:
        local = st.session_state.get("riwayat_local", {})
        local[audit_id] = data
        st.session_state["riwayat_local"] = local
        return
    try:
        ws      = get_or_create_sheet(spreadsheet, SHEET_RIWAYAT, RIWAYAT_HEADERS)
        result  = data.get("result", {})
        summary = result.get("summary", {})
        ws.append_row([
            audit_id,
            data.get("timestamp", ""),
            "|".join(data.get("files", [])),
            data.get("mode", ""),
            summary.get("overall_score", 0),
            summary.get("kritikal", 0),
            summary.get("minor", 0),
            result.get("report_type", ""),
            json.dumps(result.get("properties", []), ensure_ascii=False),
            summary.get("executive_summary", ""),
            json.dumps(result.get("findings", []), ensure_ascii=False),
        ], value_input_option="USER_ENTERED")
    except Exception as e:
        st.warning(f"⚠️ Gagal menyimpan riwayat: {e}")
        local = st.session_state.get("riwayat_local", {})
        local[audit_id] = data
        st.session_state["riwayat_local"] = local


def clear_riwayat(spreadsheet):
    if spreadsheet is None:
        st.session_state["riwayat_local"] = {}
        return
    try:
        ws = spreadsheet.worksheet(SHEET_RIWAYAT)
        ws.clear()
        ws.append_row(RIWAYAT_HEADERS)
    except Exception as e:
        st.warning(f"⚠️ Gagal menghapus riwayat: {e}")


REFERENSI_HEADERS = ["nama_laporan", "keterangan", "jumlah"]


def load_data_laporan(spreadsheet) -> dict:
    if spreadsheet is None:
        return st.session_state.get("data_laporan_local", {})
    try:
        ws      = get_or_create_sheet(spreadsheet, SHEET_REFERENSI, REFERENSI_HEADERS)
        records = ws.get_all_records()
        result  = {}
        for row in records:
            lap = str(row.get("nama_laporan", "")).strip()
            ket = str(row.get("keterangan", "")).strip()
            jml = row.get("jumlah", 0)
            if not lap:
                continue
            result.setdefault(lap, {})
            if ket:
                try:
                    result[lap][ket] = int(jml)
                except (ValueError, TypeError):
                    result[lap][ket] = 0
        return result
    except Exception as e:
        st.warning(f"⚠️ Gagal membaca referensi: {e}")
        return st.session_state.get("data_laporan_local", {})


def save_referensi_row(spreadsheet, nama_laporan: str, keterangan: str, jumlah: int):
    if spreadsheet is None:
        local = st.session_state.get("data_laporan_local", {})
        local.setdefault(nama_laporan, {})[keterangan] = jumlah
        st.session_state["data_laporan_local"] = local
        return
    try:
        ws = get_or_create_sheet(spreadsheet, SHEET_REFERENSI, REFERENSI_HEADERS)
        ws.append_row([nama_laporan, keterangan, jumlah], value_input_option="USER_ENTERED")
    except Exception as e:
        st.warning(f"⚠️ Gagal menyimpan referensi: {e}")


def delete_referensi_row(spreadsheet, nama_laporan: str, keterangan: str):
    if spreadsheet is None:
        local = st.session_state.get("data_laporan_local", {})
        if nama_laporan in local and keterangan in local[nama_laporan]:
            del local[nama_laporan][keterangan]
        st.session_state["data_laporan_local"] = local
        return
    try:
        ws       = get_or_create_sheet(spreadsheet, SHEET_REFERENSI, REFERENSI_HEADERS)
        all_vals = ws.get_all_values()
        for i, row in enumerate(all_vals[1:], start=2):
            if len(row) >= 2 and row[0] == nama_laporan and row[1] == keterangan:
                ws.delete_rows(i)
                break
    except Exception as e:
        st.warning(f"⚠️ Gagal menghapus referensi: {e}")


def add_laporan_baru(spreadsheet, nama_laporan: str):
    if spreadsheet is None:
        local = st.session_state.get("data_laporan_local", {})
        local.setdefault(nama_laporan, {})
        st.session_state["data_laporan_local"] = local
        return
    try:
        ws = get_or_create_sheet(spreadsheet, SHEET_REFERENSI, REFERENSI_HEADERS)
        ws.append_row([nama_laporan, "", ""], value_input_option="USER_ENTERED")
    except Exception as e:
        st.warning(f"⚠️ Gagal menambah laporan baru: {e}")


# ══════════════════════════════════════════════
# HELPER: EKSTRAKSI TEKS DOKUMEN
# ══════════════════════════════════════════════

def extract_text_pdf(file) -> list:
    """Ekstrak teks per halaman dari PDF."""
    pages = []
    try:
        if PDF_ENGINE == "pdfplumber":
            import pdfplumber
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    pages.append(text)
        else:
            import PyPDF2
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                pages.append(page.extract_text() or "")
    except Exception as e:
        st.warning(f"⚠️ Gagal membaca PDF '{file.name}': {e}")
    return pages


def extract_text_docx(file) -> list:
    """Ekstrak teks dari DOCX, dibagi per ~30 paragraf."""
    try:
        doc        = Document(file)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Sertakan juga teks dari tabel
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        pages, chunk = [], []
        for i, p in enumerate(paragraphs):
            chunk.append(p)
            if (i + 1) % 30 == 0:
                pages.append("\n".join(chunk))
                chunk = []
        if chunk:
            pages.append("\n".join(chunk))
        return pages
    except Exception as e:
        st.warning(f"⚠️ Gagal membaca DOCX '{file.name}': {e}")
        return []


def pages_to_text(pages: list, max_chars: int = MAX_CHARS) -> str:
    parts, total = [], 0
    for i, page in enumerate(pages, 1):
        chunk = f"\n--- Halaman {i} ---\n{page}"
        if total + len(chunk) > max_chars:
            parts.append("\n\n[... konten dipotong karena terlalu panjang ...]")
            break
        parts.append(chunk)
        total += len(chunk)
    return "".join(parts)


# ══════════════════════════════════════════════
# ENGINE: PENGECEKAN MATEMATIKA RESUME PENILAIAN
# ══════════════════════════════════════════════

def parse_id_number(s):
    if not s:
        return None
    s = str(s).strip()
    s = re.sub(r'^(Rp\.?\s*|US\$\s*|USD\s*)', '', s, flags=re.IGNORECASE).strip()
    s = s.rstrip('.,').strip()
    if re.match(r'^[\d\.]+,\d{1,2}$', s):
        s = s.replace('.', '').replace(',', '.')
    else:
        s = s.replace('.', '').replace(',', '')
    try:
        return float(s)
    except (ValueError, TypeError):
        return None


def cek_matematika_resume(doc_text: str) -> dict:
    TOLERANSI_BULAT = 0.005
    TOLERANSI_KURS  = 0.01
    findings, extracted = [], {}

    # Ekstrak kurs BI
    kurs_bi = None
    for pat in [
        r'(?:US\$\s*1\s*=\s*Rp|1\s*US\$\s*=\s*Rp)\s*([\d\.]+(?:,\d+)?)',
        r'(?:kurs|nilai tukar)[^\n]*(?:US\$|USD)[^\n]*Rp\s*([\d\.]+(?:,\d+)?)',
    ]:
        km = re.search(pat, doc_text, re.IGNORECASE)
        if km:
            v = parse_id_number(km.group(1))
            if v and v > 10000:
                kurs_bi = v
                extracted["kurs_bi"] = v
                break

    # Ekstrak komponen resume (format SRR: A. Nama ... angka_rp angka_usd)
    comp_pat = re.compile(
        r"[ \t]*([A-F])\.[ \t]+(.+?)[ \t]{5,}"
        r"([\d]{1,3}(?:\.[\d]{3})+)[ \t]+"
        r"([\d]{1,3}(?:\.[\d]{3})+)",
        re.MULTILINE
    )
    components = {}
    zone = doc_text[:60000]
    for m in comp_pat.finditer(zone):
        rp_v  = parse_id_number(m.group(3))
        usd_v = parse_id_number(m.group(4))
        if rp_v and rp_v > 0 and m.group(1) not in components:
            components[m.group(1)] = {
                "nama": m.group(2).strip()[:60],
                "rp": rp_v, "usd": usd_v,
                "rp_raw": m.group(3), "usd_raw": m.group(4),
            }
    if not components:
        comp_pat2 = re.compile(
            r"[ \t]*([A-F])\.[ \t]+(.+?)[ \t]{3,}"
            r"([\d]{1,3}(?:[.,][\d]{3})*(?:[.,][\d]{1,2})?)[ \t]+"
            r"([\d]{1,3}(?:[.,][\d]{3})*(?:[.,][\d]{1,2})?)",
            re.MULTILINE
        )
        for m in comp_pat2.finditer(zone):
            rp_v  = parse_id_number(m.group(3))
            usd_v = parse_id_number(m.group(4))
            if rp_v and rp_v > 0 and m.group(1) not in components:
                components[m.group(1)] = {
                    "nama": m.group(2).strip()[:60],
                    "rp": rp_v, "usd": usd_v,
                    "rp_raw": m.group(3), "usd_raw": m.group(4),
                }
    extracted["components"] = components

    # Ekstrak total
    total_rp = total_usd = None
    resume_ribuan = bool(re.search(
        r'Rp\s*\.?000|Rp\s+[Rr]ibu|\(Rp\s*,000|\(Rp\s+000',
        doc_text[:60000], re.IGNORECASE
    ))
    extracted["resume_ribuan"] = resume_ribuan
    for pat in [
        r'Total\s+[A-F\-]+\s+([\d]{1,3}(?:\.[\d]{3})+)[ \t]+([\d]{1,3}(?:\.[\d]{3})+)',
        r'(?:Total|TOTAL)\s+([\d]{1,3}(?:[.,][\d]{3})*(?:[.,][\d]{1,2})?)\s+([\d]{1,3}(?:[.,][\d]{3})*(?:[.,][\d]{1,2})?)',
    ]:
        tm = re.search(pat, doc_text[:60000], re.IGNORECASE)
        if tm:
            tv = parse_id_number(tm.group(1))
            if tv and tv > 0:
                total_rp  = tv
                total_usd = parse_id_number(tm.group(2)) if tm.lastindex >= 2 else None
                extracted["total_rp"]  = total_rp
                extracted["total_usd"] = total_usd
                break

    # Ekstrak nilai kesimpulan
    conclusion_rp = conclusion_usd = None
    cm = re.compile(
        r'(?:nilai pasar|NILAI PASAR|kesimpulan|KESIMPULAN)[^\n]{0,120}'
        r'Rp\s*([\d]{1,3}(?:[.,]\d{3})*(?:[.,]\d{1,2})?)',
        re.IGNORECASE | re.DOTALL
    ).search(doc_text)
    if cm:
        conclusion_rp = parse_id_number(cm.group(1))
        um = re.search(r'US\$\s*([\d]{1,3}(?:[.,]\d{3})*(?:[.,]\d{1,2})?)',
                       doc_text[cm.start():cm.start()+500], re.IGNORECASE)
        if um:
            conclusion_usd = parse_id_number(um.group(1))
    if conclusion_rp:
        extracted["conclusion_rp"]  = conclusion_rp
        extracted["conclusion_usd"] = conclusion_usd

    # Ekstrak BPB items
    bpb_items = []
    for m in re.compile(
        r'(?:Tanah|Bangunan|Sarana|Mesin|Kendaraan)[^\n]{0,80}'
        r'([\d]{1,3}(?:[.,]\d{3})*)\s+([\d,\.]+)%\s*([\d]{1,3}(?:[.,]\d{3})*)',
        re.IGNORECASE
    ).finditer(doc_text):
        bpb_v = parse_id_number(m.group(1))
        mkt_v = parse_id_number(m.group(3))
        try:
            depr_pct = float(m.group(2).replace(',', '.')) / 100
        except ValueError:
            continue
        if bpb_v and mkt_v and bpb_v > 0:
            bpb_items.append({
                "bpb": bpb_v, "depr_pct": depr_pct, "mkt_val": mkt_v,
                "expected": bpb_v * (1 - depr_pct),
            })
    extracted["bpb_items"] = bpb_items

    # Cek 1: Penjumlahan Rp
    if components and total_rp:
        sum_rp = sum(c["rp"] for c in components.values())
        diff   = abs(sum_rp - total_rp)
        ok     = diff <= total_rp * TOLERANSI_BULAT
        last_k = list(components.keys())[-1]
        findings.append({
            "id": "M001", "severity": "ok" if ok else "kritikal",
            "category": "Formula Resume",
            "title": "Penjumlahan Rp sesuai" if ok else "Penjumlahan Rp TIDAK sesuai",
            "detail": (
                f"Jumlah komponen A-{last_k}: Rp {sum_rp:,.0f}\n"
                f"Total tercantum: Rp {total_rp:,.0f}\n"
                + (f"Selisih: Rp {diff:,.0f} (dalam toleransi)" if ok
                   else f"SELISIH: Rp {diff:,.0f} — perlu diperiksa")
            ),
            "page_hint": "Resume Penilaian", "formula": f"Σ komponen Rp = {sum_rp:,.0f}",
        })

    # Cek 2: Penjumlahan USD
    usd_comps = {k: v for k, v in components.items() if v.get("usd")}
    if usd_comps and total_usd:
        sum_usd = sum(c["usd"] for c in usd_comps.values())
        diff    = abs(sum_usd - total_usd)
        ok      = diff <= total_usd * TOLERANSI_BULAT
        findings.append({
            "id": "M002", "severity": "ok" if ok else "kritikal",
            "category": "Formula Resume",
            "title": "Penjumlahan USD sesuai" if ok else "Penjumlahan USD TIDAK sesuai",
            "detail": (
                f"Jumlah USD komponen: {sum_usd:,.0f}\n"
                f"Total USD tercantum: {total_usd:,.0f}\n"
                + (f"Selisih: {diff:,.0f}" if ok else f"SELISIH: {diff:,.0f} — perlu diperiksa")
            ),
            "page_hint": "Resume Penilaian", "formula": f"Σ komponen USD = {sum_usd:,.0f}",
        })

    # Cek 3: Konversi kurs per komponen
    if kurs_bi and usd_comps:
        errors = []
        rp_mult = 1000 if resume_ribuan else 1
        for kode, comp in usd_comps.items():
            exp_usd  = (comp["rp"] * rp_mult) / kurs_bi
            diff_pct = abs(exp_usd - comp["usd"]) / exp_usd if exp_usd else 0
            if diff_pct > TOLERANSI_KURS:
                errors.append(
                    f"{kode}. {comp['nama'][:30]}: "
                    f"Rp {comp['rp']:,.0f} / {kurs_bi:,.0f} = "
                    f"USD {round(exp_usd):,.0f} (tercantum: {comp['usd']:,.0f}, "
                    f"selisih: {abs(comp['usd'] - round(exp_usd)):,.0f})"
                )
        ok = len(errors) == 0
        findings.append({
            "id": "M003", "severity": "ok" if ok else "kritikal",
            "category": "Konversi Kurs",
            "title": "Konversi kurs per komponen sesuai" if ok else f"Konversi kurs SALAH pada {len(errors)} komponen",
            "detail": (
                f"Semua {len(usd_comps)} komponen: Rp / {kurs_bi:,.0f} = USD (dalam toleransi)" if ok
                else f"Kurs BI: 1 USD = Rp {kurs_bi:,.0f}\n" + "\n".join(errors)
            ),
            "page_hint": "Resume Penilaian", "formula": f"Nilai Rp / {kurs_bi:,.0f} = Nilai USD",
        })

    # Cek 4: Konversi kurs total
    if kurs_bi and total_rp and total_usd:
        rp_mult       = 1000 if resume_ribuan else 1
        exp_total_usd = (total_rp * rp_mult) / kurs_bi
        diff_pct      = abs(exp_total_usd - total_usd) / exp_total_usd if exp_total_usd else 1
        ok   = diff_pct <= TOLERANSI_KURS
        exp_r = round(exp_total_usd)
        findings.append({
            "id": "M004", "severity": "ok" if ok else "kritikal",
            "category": "Konversi Kurs",
            "title": "Konversi kurs total sesuai" if ok else "Konversi kurs TOTAL tidak sesuai",
            "detail": (
                f"Total Rp: {total_rp:,.0f}\nKurs BI: {kurs_bi:,.0f}\n"
                f"USD seharusnya: {exp_r:,.0f}\nUSD tercantum: {total_usd:,.0f}\n"
                + ("OK" if ok else f"SELISIH: {abs(total_usd - exp_r):,.0f} ({diff_pct*100:.2f}%)")
            ),
            "page_hint": "Resume / Kesimpulan",
            "formula": f"{total_rp:,.0f} / {kurs_bi:,.0f} = {exp_r:,.0f}",
        })

    # Cek 5: Konsistensi nilai kesimpulan vs total resume
    if total_rp and conclusion_rp:
        norm_total = total_rp * 1000
        diff_pct   = abs(norm_total - conclusion_rp) / conclusion_rp if conclusion_rp else 1
        ok = diff_pct <= TOLERANSI_BULAT
        findings.append({
            "id": "M005", "severity": "ok" if ok else "kritikal",
            "category": "Konsistensi Nilai",
            "title": "Nilai kesimpulan konsisten dengan resume" if ok else "Nilai kesimpulan TIDAK konsisten",
            "detail": (
                f"Total resume (ribuan): Rp {total_rp:,.0f}\n"
                f"Setara penuh: Rp {norm_total:,.0f}\n"
                f"Nilai kesimpulan: Rp {conclusion_rp:,.0f}\n"
                + ("Konsisten" if ok else f"SELISIH: Rp {abs(norm_total - conclusion_rp):,.0f}")
            ),
            "page_hint": "Surat Pengantar / Kesimpulan",
            "formula": "Resume × 1.000 = Nilai Kesimpulan",
        })

    # Cek 6: BPB
    if bpb_items:
        errors = []
        for item in bpb_items[:20]:
            diff_pct = abs(item["expected"] - item["mkt_val"]) / item["expected"] if item["expected"] > 0 else 0
            if diff_pct > TOLERANSI_BULAT and abs(item["expected"] - item["mkt_val"]) > 1000:
                errors.append(
                    f"BPB {item['bpb']:,.0f} × (1-{item['depr_pct']*100:.1f}%) = "
                    f"{round(item['expected']):,.0f} (tercantum: {item['mkt_val']:,.0f}, "
                    f"selisih: {abs(item['mkt_val'] - round(item['expected'])):,.0f})"
                )
        ok = len(errors) == 0
        findings.append({
            "id": "M006", "severity": "ok" if ok else "minor",
            "category": "Formula BPB",
            "title": f"Formula BPB sesuai — {len(bpb_items)} baris" if ok else f"Formula BPB perlu verifikasi — {len(errors)} baris",
            "detail": (
                f"Semua {len(bpb_items)} baris BPB memenuhi formula" if ok
                else "Potensi ketidaksesuaian:\n" + "\n".join(errors[:5])
            ),
            "page_hint": "Bab D.1 / Tabel Ringkasan",
            "formula": "BPB × (1 - Depresiasi%) = Nilai Pasar",
        })

    if not findings:
        findings.append({
            "id": "M000", "severity": "info",
            "category": "Formula Resume",
            "title": "Data numerik tidak dapat diekstrak otomatis",
            "detail": (
                "Engine tidak menemukan tabel resume dengan format yang dapat diparse. "
                "Kemungkinan: (1) tidak ada Resume Penilaian, (2) format tabel non-standar, "
                "atau (3) PDF ter-scan."
            ),
            "page_hint": "Resume Penilaian", "formula": "-",
        })

    return {
        "findings":   findings,
        "extracted":  extracted,
        "n_ok":       sum(1 for f in findings if f["severity"] == "ok"),
        "n_kritikal": sum(1 for f in findings if f["severity"] == "kritikal"),
        "n_minor":    sum(1 for f in findings if f["severity"] == "minor"),
    }


# ══════════════════════════════════════════════
# HELPER: CALL CLAUDE API
# ══════════════════════════════════════════════

def recover_partial_json(raw_text: str):
    start = raw_text.find("{")
    if start == -1:
        return None
    partial = raw_text[start:]

    def close_json(text):
        cleaned = re.sub(r',\s*"[^"]*"\s*:?\s*$', "", text.rstrip())
        cleaned = re.sub(r',\s*"[^"]*"\s*$', "", cleaned.rstrip())
        stack, in_string, esc = [], False, False
        for ch in cleaned:
            if esc: esc = False; continue
            if ch == "\\" and in_string: esc = True; continue
            if ch == '"': in_string = not in_string; continue
            if in_string: continue
            if ch in "{[": stack.append(ch)
            elif ch in "}]":
                if stack: stack.pop()
        return cleaned + "".join("]" if b == "[" else "}" for b in reversed(stack))

    try:
        parsed = json.loads(close_json(partial))
        parsed.setdefault("summary", {})
        parsed.setdefault("findings", [])
        parsed.setdefault("properties", [])
        parsed["_partial"] = True
        return parsed
    except Exception:
        pass

    result = {"_partial": True, "findings": [], "properties": [], "summary": {}}
    for field in ["report_type"]:
        m = re.search(f'"{field}"\\s*:\\s*"([^"]*)"', partial)
        if m: result[field] = m.group(1)
    for field in ["total_findings", "kritikal", "minor", "ok", "info", "overall_score"]:
        nm = re.search(f'"{field}"\\s*:\\s*(\\d+)', partial)
        if nm: result["summary"][field] = int(nm.group(1))
    em = re.search(r'"executive_summary"\s*:\s*"([^"]*)"', partial)
    if em: result["summary"]["executive_summary"] = em.group(1)
    return result if result.get("report_type") or result["summary"] else None


def call_claude(api_key: str, mode_instruction: str, check_items: list,
                doc_text: str, excel_context: str = "") -> tuple:
    """
    Panggil Claude API dengan instruksi mode, item cek, teks dokumen, dan (opsional) data Excel.
    Returns (parsed_json, raw_text)
    """
    client = anthropic.Anthropic(api_key=api_key)

    excel_section = ""
    if excel_context:
        excel_section = f"\n\n{excel_context}"

    user_message = (
        f"{mode_instruction}\n\n"
        f"Item yang harus diperiksa:\n"
        + "\n".join(f"- {item}" for item in check_items)
        + excel_section
        + f"\n\nKONTEN DOKUMEN:\n{doc_text}"
    )

    response = client.messages.create(
        model=MODEL,
        max_tokens=MAX_TOKENS,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_message}],
    )
    raw_text = response.content[0].text

    def try_parse(text):
        try:
            return json.loads(text)
        except Exception:
            return None

    parsed = try_parse(raw_text)
    if parsed is None:
        stripped = re.sub(r"```(?:json)?\s*([\s\S]*?)```", r"\1", raw_text).strip()
        parsed = try_parse(stripped)
    if parsed is None:
        for m in re.finditer(r"\{[\s\S]*\}", raw_text):
            parsed = try_parse(m.group())
            if parsed: break
    if parsed is None:
        fixed = re.sub(r",\s*([}\]])", r"\1", raw_text)
        parsed = try_parse(fixed)
    if parsed is None:
        parsed = recover_partial_json(raw_text)
    if parsed is None:
        raise ValueError(
            f"Tidak bisa mem-parse JSON dari response Claude.\n"
            f"Raw (500 char pertama):\n{raw_text[:500]}"
        )
    return parsed, raw_text


# ══════════════════════════════════════════════
# HELPER: GENERATE LAPORAN REVIEW (DOWNLOAD)
# ══════════════════════════════════════════════

def generate_download_report(result: dict, math_result: dict,
                              local_result: dict, files: list,
                              mode_label: str) -> str:
    """Buat teks laporan review untuk diunduh."""
    lines = []
    now   = datetime.now().strftime("%d %B %Y, %H:%M")
    summary = result.get("summary", {})
    lines += [
        "=" * 70,
        "LAPORAN REVIEW DOKUMEN",
        f"CekLaporan v7.0 — KJPP SRR",
        f"Tanggal  : {now}",
        f"File     : {', '.join(files)}",
        f"Mode     : {mode_label}",
        f"Skor QC  : {summary.get('overall_score', 0)}/100",
        "=" * 70, "",
        "RINGKASAN EKSEKUTIF",
        "-" * 70,
        summary.get("executive_summary", "-"), "",
        f"Total temuan : {summary.get('total_findings', 0)}",
        f"🔴 Kritikal   : {summary.get('kritikal', 0)}",
        f"🟡 Minor      : {summary.get('minor', 0)}",
        f"🟢 Sesuai     : {summary.get('ok', 0)}",
        "", "=" * 70,
        "DETAIL TEMUAN (AI AUDIT)",
        "=" * 70,
    ]

    for sev in ["kritikal", "minor", "ok", "info"]:
        group = [f for f in result.get("findings", []) if f.get("severity") == sev]
        if not group:
            continue
        lines.append(f"\n{'─'*60}")
        lines.append(f"[{sev.upper()}] — {len(group)} temuan")
        lines.append(f"{'─'*60}")
        for f in group:
            lines += [
                f"",
                f"[{f.get('id','')}] {f.get('title','')}",
                f"  Kategori  : {f.get('category','')}",
                f"  Lokasi    : {f.get('page_hint','')}",
                f"  Detail    :",
            ]
            for dl in f.get("detail", "").split("\n"):
                lines.append(f"    {dl}")

    lines += ["", "=" * 70, "PENGECEKAN LOKAL (OTOMATIS)", "=" * 70]
    for f in local_result.get("findings", []):
        lines += [
            f"",
            f"[{f.get('id','')}] [{f.get('severity','').upper()}] {f.get('title','')}",
            f"  {f.get('detail','')}",
        ]

    lines += ["", "=" * 70, "PENGECEKAN MATEMATIKA", "=" * 70]
    for f in math_result.get("findings", []):
        lines += [
            f"",
            f"[{f.get('id','')}] [{f.get('severity','').upper()}] {f.get('title','')}",
            f"  Formula : {f.get('formula','-')}",
            f"  Detail  : {f.get('detail','')}",
        ]

    lines += [
        "", "=" * 70,
        f"CekLaporan v7.0 · KJPP SRR · Powered by Claude AI ({MODEL})",
        "=" * 70,
    ]
    return "\n".join(lines)


# ══════════════════════════════════════════════
# HELPER: RENDER UI
# ══════════════════════════════════════════════

def render_finding_card(f: dict):
    sev   = f.get("severity", "info")
    cfg   = SEVERITY_CONFIG.get(sev, SEVERITY_CONFIG["info"])
    prop  = f.get("property", "")
    prop_tag = f' &nbsp;·&nbsp; <span style="color:#1e6fbf;">📌 {prop}</span>' if prop else ""
    st.markdown(f"""
<div style="background:{cfg['bg']};border:1px solid {cfg['color']}40;border-left:4px solid {cfg['color']};
            border-radius:8px;padding:14px 16px;margin-bottom:10px;">
    <div style="display:flex;align-items:center;gap:8px;margin-bottom:6px;">
        <span style="background:{cfg['color']}22;color:{cfg['color']};border:1px solid {cfg['color']};
                     font-size:10px;font-weight:700;padding:2px 8px;border-radius:4px;
                     text-transform:uppercase;">{cfg['emoji']} {sev}</span>
        <span style="color:#666;font-size:11px;font-family:monospace;">{f.get('category','')}{prop_tag}</span>
        <span style="color:#666;font-size:11px;font-family:monospace;margin-left:auto;">
            {f.get('id','')} &nbsp; {f.get('page_hint','')}</span>
    </div>
    <div style="font-size:14px;font-weight:600;color:#1a1a2e;margin-bottom:6px;">{f.get('title','')}</div>
    <div style="font-size:12px;color:#444;font-family:monospace;background:#ffffff;
                padding:8px 12px;border-radius:6px;line-height:1.6;border:1px solid #e0e0e0;
                white-space:pre-line;">{f.get('detail','')}</div>
</div>""", unsafe_allow_html=True)


def render_summary_cards(s: dict):
    score  = s.get("overall_score", 0)
    sc     = "#1a9e67" if score >= 80 else "#d4860a" if score >= 60 else "#c0392b"
    kritis = s.get("kritis", 0) + s.get("kritikal", 0)
    mayor  = s.get("mayor", 0)
    minor  = s.get("minor", 0)
    status = s.get("status", "")
    status_color = "#1a9e67" if "LULUS" in status else "#e67e22" if "MINOR" in status else "#c0392b"

    # Status banner
    if status:
        st.markdown(f"""
<div style="background:{status_color}11;border:1px solid {status_color}44;border-radius:8px;
            padding:8px 16px;margin-bottom:12px;text-align:center;">
    <span style="font-size:13px;font-weight:700;color:{status_color};">{status}</span>
</div>""", unsafe_allow_html=True)

    c1, c2, c3, c4, c5 = st.columns(5)
    for col, num, label, color, bg in [
        (c1, score,  "Skor QC",     sc,        "#f8fafb"),
        (c2, kritis, "🔴 Kritis",   "#c0392b", "#fff0f0"),
        (c3, mayor,  "🟠 Mayor",    "#e67e22", "#fff5ec"),
        (c4, minor,  "🟡 Minor",    "#d4860a", "#fff8e6"),
        (c5, s.get("ok", 0), "🟢 Sesuai", "#1a9e67", "#edfaf4"),
    ]:
        col.markdown(f"""
<div style="background:{bg};border:1px solid #dde3ea;border-radius:10px;padding:14px 8px;
            text-align:center;box-shadow:0 1px 4px rgba(0,0,0,0.06);">
    <div style="font-size:28px;font-weight:800;color:{color};font-family:monospace;">{num}</div>
    <div style="font-size:10px;color:#6b7280;text-transform:uppercase;letter-spacing:1px;">{label}</div>
</div>""", unsafe_allow_html=True)


def extract_keywords(finding: dict) -> list:
    keywords, detail, title = [], finding.get("detail",""), finding.get("title","")
    for pat in [r"Rp\s*([\d][.\d,]+)", r"US\$\s*([\d][.\d,]+)",
                r"([\d]{3}\.[\d]{3}\.[\d]{3})"]:
        for m in re.findall(pat, detail+" "+title, re.IGNORECASE):
            cleaned = m.strip().rstrip(".,")
            if cleaned: keywords.append(cleaned)
    for pat in [
        r"\d{1,2}\s+(?:Januari|Februari|Maret|April|Mei|Juni|Juli|Agustus|September|Oktober|November|Desember)\s+\d{4}",
        r"No\.\s+[\w\-/]+", r"SHGB?\s+No[\.\s]+\d+",
    ]:
        for m in re.findall(pat, detail+" "+title, re.IGNORECASE):
            if len(m) > 4: keywords.append(m.strip())
    seen, result = set(), []
    for kw in sorted(keywords, key=len, reverse=True):
        norm = kw.lower().replace(" ","")
        if norm not in seen and len(kw) >= 3:
            seen.add(norm)
            result.append(kw)
    return result[:8]


def parse_page_hints(page_hint: str) -> list:
    if not page_hint: return []
    romawi_map = {"i":1,"ii":2,"iii":3,"iv":4,"v":5,"vi":6,"vii":7,
                  "viii":8,"ix":9,"x":10,"xi":11,"xii":12,"xiii":13,
                  "xiv":14,"xv":15,"xvi":16,"xvii":17,"xviii":18,"xix":19,"xx":20}
    pages = []
    for m in re.finditer(r'(?:hal|halaman|page)[\.\s]+([^\s,;()vs]+)', page_hint, re.IGNORECASE):
        raw = m.group(1).strip().lower().rstrip(".,)")
        if raw in romawi_map: pages.append(str(romawi_map[raw]))
        elif raw.isdigit(): pages.append(raw)
        else:
            rng = re.match(r'(\d+)[-–](\d+)', raw)
            if rng:
                s, e = int(rng.group(1)), int(rng.group(2))
                pages.extend([str(i) for i in range(s, min(e+1, s+15))])
    for rn, rv in romawi_map.items():
        if re.search(r'\b' + rn + r'\b', page_hint, re.IGNORECASE):
            val = str(rv)
            if val not in pages: pages.append(val)
    return list(dict.fromkeys(pages))[:12]


def find_pages_for_finding(finding: dict, all_pages: list) -> list:
    if not all_pages: return []
    keywords   = extract_keywords(finding)
    hint_pages = parse_page_hints(finding.get("page_hint", ""))
    results    = []
    romawi_map = {"i":1,"ii":2,"iii":3,"iv":4,"v":5,"vi":6,"vii":7,
                  "viii":8,"ix":9,"x":10,"xi":11,"xii":12,"xiii":13,
                  "xiv":14,"xv":15,"xvi":16,"xvii":17,"xviii":18,"xix":19,"xx":20}
    for page_num, page_text in enumerate(all_pages, start=1):
        if not page_text or not page_text.strip(): continue
        score = 0
        matched_kws = []
        for hp in hint_pages:
            if hp == str(page_num): score += 30
            elif hp.lower() in romawi_map and romawi_map[hp.lower()] == page_num: score += 30
        page_lower = page_text.lower()
        for kw in keywords:
            try:
                count = len(re.findall(re.escape(kw.lower()), page_lower))
            except re.error:
                count = page_lower.count(kw.lower())
            if count > 0:
                score += min(count * 10, 30)
                matched_kws.append(kw)
        if score > 0:
            results.append({
                "page_num": page_num, "text": page_text,
                "matches": matched_kws, "score": score,
                "hint_match": any(hp == str(page_num) or
                    (hp.lower() in romawi_map and romawi_map[hp.lower()] == page_num)
                    for hp in hint_pages),
            })
    results.sort(key=lambda x: (x["hint_match"], x["score"]), reverse=True)
    return results[:5]


def highlight_keywords_in_text(text: str, keywords: list) -> str:
    escaped = text.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
    for kw in sorted(keywords, key=len, reverse=True):
        if not kw or len(kw) < 3: continue
        try:
            escaped = re.compile(re.escape(kw), re.IGNORECASE).sub(
                lambda m: f'<mark style="background:#fef08a;color:#1a1a2e;'
                          f'padding:0 2px;border-radius:2px;font-weight:700;">{m.group()}</mark>',
                escaped
            )
        except re.error:
            pass
    return escaped


def render_finding_with_preview(finding: dict, all_pages: list, card_key: str):
    sev   = finding.get("severity", "info")
    cfg   = SEVERITY_CONFIG.get(sev, SEVERITY_CONFIG["info"])
    prop  = finding.get("property", "")
    prop_tag = f' &nbsp;·&nbsp; <span style="color:#1e6fbf;">📌 {prop}</span>' if prop else ""
    st.markdown(f"""
<div style="background:{cfg['bg']};border:1px solid {cfg['color']}40;border-left:4px solid {cfg['color']};
            border-radius:8px;padding:14px 16px;margin-bottom:4px;">
    <div style="display:flex;align-items:center;gap:8px;margin-bottom:6px;">
        <span style="background:{cfg['color']}22;color:{cfg['color']};border:1px solid {cfg['color']};
                     font-size:10px;font-weight:700;padding:2px 8px;border-radius:4px;
                     text-transform:uppercase;">{cfg['emoji']} {sev}</span>
        <span style="color:#666;font-size:11px;font-family:monospace;">
            {finding.get('category','')}{prop_tag}</span>
        <span style="color:#666;font-size:11px;font-family:monospace;margin-left:auto;">
            {finding.get('id','')} &nbsp; {finding.get('page_hint','')}</span>
    </div>
    <div style="font-size:14px;font-weight:600;color:#1a1a2e;margin-bottom:6px;">{finding.get('title','')}</div>
    <div style="font-size:12px;color:#444;font-family:monospace;background:#ffffff;
                padding:8px 12px;border-radius:6px;line-height:1.6;border:1px solid #e0e0e0;
                white-space:pre-line;">{finding.get('detail','')}</div>
</div>""", unsafe_allow_html=True)

    if all_pages:
        keywords      = extract_keywords(finding)
        btn_label     = f"📄 Lihat Preview Dokumen"
        if finding.get("page_hint"):
            btn_label += f" ({finding['page_hint']})"
        with st.expander(btn_label, expanded=False):
            relevant_pages = find_pages_for_finding(finding, all_pages)
            if not relevant_pages:
                st.caption("⚠️ Tidak ada halaman yang cocok ditemukan.")
            else:
                st.caption(
                    f"Ditemukan **{len(relevant_pages)}** halaman relevan · "
                    f"Keywords: {', '.join(f'`{k}`' for k in keywords[:4])}"
                )
                for pr in relevant_pages:
                    badge   = "📍 sesuai hint" if pr["hint_match"] else f"skor: {pr['score']}"
                    kw_str  = ", ".join(f'`{k}`' for k in pr["matches"][:4]) if pr["matches"] else "—"
                    st.markdown(
                        f'<div style="display:flex;align-items:center;gap:8px;margin:8px 0 4px;">'
                        f'<span style="background:#1a9e67;color:#fff;font-size:10px;font-weight:700;'
                        f'padding:2px 8px;border-radius:4px;font-family:monospace;">HAL. {pr["page_num"]}</span>'
                        f'<span style="font-size:11px;color:#6b7280;font-family:monospace;">{badge}</span>'
                        f'<span style="font-size:11px;color:#6b7280;">Keywords: {kw_str}</span>'
                        f'</div>', unsafe_allow_html=True
                    )
                    page_text = pr["text"]
                    if len(page_text) > 1500:
                        best_pos = 0
                        for kw in pr["matches"][:3]:
                            pos = page_text.lower().find(kw.lower())
                            if pos > 0:
                                best_pos = max(0, pos - 200)
                                break
                        page_text = ("…" if best_pos > 0 else "") + page_text[best_pos:best_pos+1500]
                        if best_pos + 1500 < len(pr["text"]):
                            page_text += "…"
                    highlighted = highlight_keywords_in_text(page_text, keywords)
                    st.markdown(
                        f'<div style="background:#f8fafb;border:1px solid #dde3ea;border-radius:6px;'
                        f'padding:12px 14px;font-family:monospace;font-size:12px;line-height:1.7;'
                        f'color:#374151;white-space:pre-wrap;word-break:break-word;max-height:350px;'
                        f'overflow-y:auto;">{highlighted}</div>', unsafe_allow_html=True
                    )
                    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("<div style='margin-bottom:8px;'></div>", unsafe_allow_html=True)


def render_findings_as_sections(findings: list, all_pages: list):
    """
    Render findings grouped into numbered sections with HTML table format,
    matching the design: section header + table per category group.
    """
    if not findings:
        st.success("✅ Tidak ada temuan AI — laporan terlihat konsisten.")
        return

    # Helper HTML for the standard bordered table
    TABLE_STYLE = (
        "width:100%;border-collapse:collapse;font-size:13px;"
        "background:#fff;border:1px solid #dde3ea;border-radius:8px;"
        "overflow:hidden;margin-bottom:24px;"
    )
    TH_STYLE = "padding:10px 12px;text-align:left;font-size:11px;color:#6b7280;background:#f8fafb;border-bottom:2px solid #dde3ea;border-right:1px solid #dde3ea;"
    TD_STYLE = "padding:10px 12px;border-bottom:1px solid #f0f0f0;border-right:1px solid #f0f0f0;vertical-align:top;"
    TD_NO_STYLE = "padding:10px 12px;border-bottom:1px solid #f0f0f0;border-right:1px solid #f0f0f0;vertical-align:top;width:36px;color:#9ca3af;font-size:12px;font-family:monospace;"
    TD_LOKASI_STYLE = "padding:10px 12px;border-bottom:1px solid #f0f0f0;border-right:1px solid #f0f0f0;vertical-align:top;width:160px;font-size:12px;color:#6b7280;font-family:monospace;"

    def severity_badge(sev):
        cfg = SEVERITY_CONFIG.get(sev, SEVERITY_CONFIG["info"])
        return (f'<span style="background:{cfg["color"]}22;color:{cfg["color"]};'
                f'border:1px solid {cfg["color"]};font-size:10px;font-weight:700;'
                f'padding:1px 6px;border-radius:3px;text-transform:uppercase;">'
                f'{cfg["emoji"]} {sev}</span>')

    def bold_text(text):
        """Bold words between ** markers."""
        return re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', str(text))

    # Categorize findings
    section_findings = {s["key"]: [] for s in SECTION_CONFIG}
    catch_all_key = SECTION_CONFIG[-1]["key"]  # "lainlain"

    category_to_key = {}
    for s in SECTION_CONFIG:
        for cat in s["categories"]:
            category_to_key[cat.lower()] = s["key"]

    for f in findings:
        cat = f.get("category", "").strip()
        skey = category_to_key.get(cat.lower())
        if skey is None:
            skey = catch_all_key
        section_findings[skey].append(f)

    # Render each section that has findings
    for sec in SECTION_CONFIG:
        items = section_findings[sec["key"]]
        if not items:
            continue

        # Section header
        n = sec['num']
        header_label = f"BAGIAN {n} — {sec['title']}" if n != "—" else sec['title']
        # Count kritis/mayor/minor in this section
        sev_counts = {}
        for f in items:
            sv = f.get("severity", "info")
            if sv in ("kritis", "kritikal"): sv = "kritis"
            sev_counts[sv] = sev_counts.get(sv, 0) + 1
        pills = ""
        for sv, clr in [("kritis","#c0392b"),("mayor","#e67e22"),("minor","#d4860a")]:
            if sev_counts.get(sv, 0):
                pills += (f'<span style="background:{clr}22;color:{clr};border:1px solid {clr};'
                          f'font-size:10px;font-weight:700;padding:1px 8px;border-radius:10px;'
                          f'margin-left:8px;">{SEVERITY_CONFIG[sv]["emoji"]} {sev_counts[sv]}</span>')

        st.markdown(f"""
<div style="margin-top:32px;margin-bottom:10px;display:flex;align-items:center;gap:0;
            border-bottom:2px solid #1a1a2e;padding-bottom:6px;">
    <span style="font-size:14px;font-weight:800;color:#1a1a2e;letter-spacing:-0.2px;
                 text-transform:uppercase;">{header_label}</span>
    {pills}
</div>""", unsafe_allow_html=True)

        style = sec["style"]

        # ── Narrative style (for Artefak Copy-Paste / Bagian B) ──
        if style == "narrative":
            # Show warning box
            st.markdown("""
<div style="background:#fff8e6;border:1px solid #f5e0a0;border-radius:8px;
            padding:10px 14px;margin-bottom:14px;font-size:13px;color:#7c5800;">
    ⚠️ Ini adalah masalah paling serius dalam dokumen ini.
</div>""", unsafe_allow_html=True)
            for idx, f in enumerate(items, 1):
                sev = f.get("severity", "info")
                cfg = SEVERITY_CONFIG.get(sev, SEVERITY_CONFIG["info"])
                sub_label = chr(96 + idx)  # 'a', 'b', 'c', ...
                title = f.get("title", "")
                detail = f.get("detail", "")
                page_hint = f.get("page_hint", "")

                # Extract quoted text from detail
                detail_lines = detail.split("\n")
                main_lines = []
                quote_lines = []
                for line in detail_lines:
                    stripped = line.strip()
                    if stripped.startswith('"') and stripped.endswith('"'):
                        quote_lines.append(stripped)
                    elif stripped.startswith('•') or stripped.startswith('-'):
                        main_lines.append(f'<li style="margin-bottom:4px;">{stripped[1:].strip()}</li>')
                    else:
                        main_lines.append(f'<p style="margin:4px 0;">{stripped}</p>')

                quotes_html = ""
                for q in quote_lines:
                    quotes_html += (
                        f'<blockquote style="border-left:3px solid #d4860a;padding:6px 12px;'
                        f'margin:8px 0;background:#fffbf0;color:#5c4000;font-style:italic;'
                        f'font-size:13px;">{q}</blockquote>'
                    )

                bullets_exist = any("<li" in l for l in main_lines)
                if bullets_exist:
                    li_items = "".join(l for l in main_lines if "<li" in l)
                    p_items = "".join(l for l in main_lines if "<li" not in l)
                    content_html = f'<ul style="padding-left:20px;margin:6px 0;">{li_items}</ul>{p_items}'
                else:
                    content_html = "".join(main_lines)

                st.markdown(f"""
<div style="background:{cfg['bg']};border:1px solid {cfg['color']}40;
            border-left:4px solid {cfg['color']};border-radius:8px;
            padding:14px 16px;margin-bottom:12px;">
    <div style="display:flex;align-items:center;gap:8px;margin-bottom:8px;">
        <span style="font-size:14px;font-weight:800;color:#1a1a2e;">
            {sec['num']}{sub_label}. {title}
        </span>
        {severity_badge(sev)}
        <span style="font-size:11px;color:#9ca3af;font-family:monospace;margin-left:auto;">{page_hint}</span>
    </div>
    <div style="font-size:13px;color:#374151;line-height:1.7;">{content_html}</div>
    {quotes_html}
</div>""", unsafe_allow_html=True)
            continue

        # ── Table styles ──
        if style == "table_angka":
            # Bagian A: No | Lokasi | Masalah | Teks Salah | Seharusnya
            headers = ["No", "Lokasi", "Masalah", "Ditemukan", "Seharusnya"]
            rows_html = ""
            for i, f in enumerate(items, 1):
                sev = f.get("severity", "mayor")
                if sev in ("kritis", "kritikal"): sev = "kritis"
                badge = severity_badge(sev)
                lokasi = f.get("page_hint", "—")
                masalah = f.get("title", "—")
                teks_salah = f.get("teks_salah") or ""
                seharusnya = f.get("seharusnya") or ""
                if not teks_salah and not seharusnya:
                    teks_salah = f.get("detail", "")[:200]
                rows_html += (
                    f'<tr>'
                    f'<td style="{TD_NO_STYLE}">{i}</td>'
                    f'<td style="{TD_LOKASI_STYLE}">{lokasi} {badge}</td>'
                    f'<td style="{TD_STYLE};font-weight:600;">{bold_text(masalah)}</td>'
                    f'<td style="{TD_STYLE}"><code style="font-size:12px;background:#fff5f5;'
                    f'padding:1px 4px;border-radius:3px;color:#c0392b;">{teks_salah}</code></td>'
                    f'<td style="{TD_STYLE}"><code style="font-size:12px;background:#f0fff4;'
                    f'padding:1px 4px;border-radius:3px;color:#1a9e67;">{seharusnya}</code></td>'
                    f'</tr>'
                )
            st.markdown(f"""
<table style="{TABLE_STYLE}">
<thead><tr>
{''.join(f'<th style="{TH_STYLE}">{h}</th>' for h in headers)}
</tr></thead>
<tbody>{rows_html}</tbody>
</table>""", unsafe_allow_html=True)

        elif style == "table_typo":
            headers = ["No", "Lokasi", "Teks Salah", "Seharusnya", "Keterangan"]
            rows_html = ""
            for i, f in enumerate(items, 1):
                sev = f.get("severity", "minor")
                lokasi = f.get("page_hint", "—")
                teks_salah = f.get("teks_salah") or f.get("title", "")
                seharusnya = f.get("seharusnya") or "—"
                # Keterangan: ambil bagian sebelum " — konteks:" dari detail
                raw_detail = f.get("detail", "")
                keterangan = raw_detail.split(" — konteks:")[0].strip() if " — konteks:" in raw_detail else raw_detail[:100]
                rows_html += (
                    f'<tr>'
                    f'<td style="{TD_NO_STYLE}">{i}</td>'
                    f'<td style="{TD_LOKASI_STYLE}">{lokasi}</td>'
                    f'<td style="{TD_STYLE}"><code style="font-size:12px;background:#fff5f5;'
                    f'padding:1px 4px;border-radius:3px;color:#c0392b;">{teks_salah}</code></td>'
                    f'<td style="{TD_STYLE}"><code style="font-size:12px;background:#f0fff4;'
                    f'padding:1px 4px;border-radius:3px;color:#1a9e67;">{seharusnya}</code></td>'
                    f'<td style="{TD_STYLE};color:#6b7280;font-size:12px;">{keterangan}</td>'
                    f'</tr>'
                )
            st.markdown(f"""
<table style="{TABLE_STYLE}">
<thead><tr>
{''.join(f'<th style="{TH_STYLE}">{h}</th>' for h in headers)}
</tr></thead>
<tbody>{rows_html}</tbody>
</table>""", unsafe_allow_html=True)

        elif style == "table_nama":
            headers = ["No", "Lokasi", "Teks", "Catatan"]
            rows_html = ""
            for i, f in enumerate(items, 1):
                sev = f.get("severity", "minor")
                badge = severity_badge(sev)
                lokasi = f.get("page_hint", "—")
                teks = f.get("title", "—")
                catatan = f.get("detail", "—")
                if len(catatan) > 200:
                    catatan = catatan[:200] + "…"
                rows_html += (
                    f'<tr>'
                    f'<td style="{TD_NO_STYLE}">{i}</td>'
                    f'<td style="{TD_LOKASI_STYLE}">{lokasi} {badge}</td>'
                    f'<td style="{TD_STYLE}"><strong>{teks}</strong></td>'
                    f'<td style="{TD_STYLE}">{bold_text(catatan)}</td>'
                    f'</tr>'
                )
            st.markdown(f"""
<table style="{TABLE_STYLE}">
<thead><tr>
{''.join(f'<th style="{TH_STYLE}">{h}</th>' for h in headers)}
</tr></thead>
<tbody>{rows_html}</tbody>
</table>""", unsafe_allow_html=True)

        elif style == "table_ejaan":
            headers = ["No", "Masalah", "Contoh"]
            rows_html = ""
            for i, f in enumerate(items, 1):
                sev = f.get("severity", "minor")
                badge = severity_badge(sev)
                masalah = f.get("title", "—")
                contoh = f.get("detail", "—")
                if len(contoh) > 300:
                    contoh = contoh[:300] + "…"
                rows_html += (
                    f'<tr>'
                    f'<td style="{TD_NO_STYLE}">{i} {badge}</td>'
                    f'<td style="{TD_STYLE}">{bold_text(masalah)}</td>'
                    f'<td style="{TD_STYLE}">{bold_text(contoh)}</td>'
                    f'</tr>'
                )
            st.markdown(f"""
<table style="{TABLE_STYLE}">
<thead><tr>
{''.join(f'<th style="{TH_STYLE}">{h}</th>' for h in headers)}
</tr></thead>
<tbody>{rows_html}</tbody>
</table>""", unsafe_allow_html=True)

        elif style == "table_lokasi_keterangan":
            headers = ["No", "Lokasi", "Keterangan"]
            rows_html = ""
            for i, f in enumerate(items, 1):
                sev = f.get("severity", "kritikal")
                badge = severity_badge(sev)
                lokasi = f.get("page_hint", "—")
                keterangan = f.get("detail") or f.get("title", "—")
                if len(keterangan) > 300:
                    keterangan = keterangan[:300] + "…"
                rows_html += (
                    f'<tr>'
                    f'<td style="{TD_NO_STYLE}">{i}</td>'
                    f'<td style="{TD_LOKASI_STYLE}">{lokasi} {badge}</td>'
                    f'<td style="{TD_STYLE}">{bold_text(keterangan)}</td>'
                    f'</tr>'
                )
            st.markdown(f"""
<table style="{TABLE_STYLE}">
<thead><tr>
{''.join(f'<th style="{TH_STYLE}">{h}</th>' for h in headers)}
</tr></thead>
<tbody>{rows_html}</tbody>
</table>""", unsafe_allow_html=True)

        else:  # table_lokasi_masalah (default)
            headers = ["No", "Lokasi", "Masalah"]
            rows_html = ""
            for i, f in enumerate(items, 1):
                sev = f.get("severity", "info")
                badge = severity_badge(sev)
                lokasi = f.get("page_hint", "—")
                masalah = f.get("detail") or f.get("title", "—")
                if len(masalah) > 400:
                    masalah = masalah[:400] + "…"
                rows_html += (
                    f'<tr>'
                    f'<td style="{TD_NO_STYLE}">{i}</td>'
                    f'<td style="{TD_LOKASI_STYLE}">{lokasi} {badge}</td>'
                    f'<td style="{TD_STYLE}">{bold_text(masalah)}</td>'
                    f'</tr>'
                )
            st.markdown(f"""
<table style="{TABLE_STYLE}">
<thead><tr>
{''.join(f'<th style="{TH_STYLE}">{h}</th>' for h in headers)}
</tr></thead>
<tbody>{rows_html}</tbody>
</table>""", unsafe_allow_html=True)


def render_excel_comparison(comparison: dict):
    """Render hasil komparasi dokumen vs Excel."""
    if not comparison:
        return
    st.markdown(f"""
<div style="background:#f5f3ff;border:1px solid #c4b5fd;border-radius:10px;
            padding:14px 16px;margin-bottom:16px;font-family:monospace;font-size:13px;">
    <b>📊 Hasil Komparasi Angka</b><br>
    <span style="color:#374151;">{comparison.get('summary','')}</span>
</div>""", unsafe_allow_html=True)

    not_found = comparison.get("not_found", [])
    if not_found:
        with st.expander(f"❓ {len(not_found)} angka di laporan TIDAK ditemukan di Excel", expanded=True):
            st.caption("Angka-angka ini ada di laporan tetapi tidak ada padanannya di lembar kerja:")
            for nf in not_found[:30]:
                ctx = nf["doc_context"][:120] + ("…" if len(nf["doc_context"]) > 120 else "")
                st.markdown(f"""
<div style="background:#fff;border:1px solid #7c3aed40;border-left:3px solid #7c3aed;
            border-radius:6px;padding:8px 12px;margin-bottom:6px;font-family:monospace;font-size:12px;">
    <span style="color:#7c3aed;font-weight:700;">{nf['doc_raw']}</span>
    &nbsp;→&nbsp; <span style="color:#6b7280;">{ctx}</span>
</div>""", unsafe_allow_html=True)

    matches = comparison.get("matches", [])
    if matches:
        with st.expander(f"✅ {len(matches)} angka cocok antara laporan dan Excel"):
            for m in matches[:20]:
                st.markdown(
                    f'<div style="font-family:monospace;font-size:11px;color:#374151;'
                    f'padding:4px 0;border-bottom:1px solid #f0f0f0;">'
                    f'<span style="color:#1a9e67;font-weight:700;">{m["doc_raw"]}</span>'
                    f' = <span style="color:#1e6fbf;">{m["excel_label"]}</span>'
                    f' [{m["excel_sheet"]} {m["excel_cell"]}]'
                    f'</div>', unsafe_allow_html=True
                )


def render_math_section(math_result: dict):
    """Render panel pengecekan matematika."""
    ext   = math_result.get("extracted", {})
    mf    = math_result.get("findings", [])
    n_comp = len(ext.get("components", {}))
    kurs_str  = f"1 USD = Rp {ext['kurs_bi']:,.0f}" if ext.get("kurs_bi") else "tidak terdeteksi"
    total_str = f"Rp {ext['total_rp']:,.0f}" if ext.get("total_rp") else "tidak terdeteksi"

    st.markdown(f"""
<div style="background:#f8fafb;border:1px solid #dde3ea;border-radius:8px;
            padding:10px 14px;margin-bottom:14px;font-family:monospace;font-size:12px;
            display:flex;gap:24px;flex-wrap:wrap;">
    <span>🏷️ <b>Komponen</b>: {n_comp}</span>
    <span>💱 <b>Kurs BI</b>: {kurs_str}</span>
    <span>📊 <b>Total Resume</b>: {total_str}</span>
    <span>✅ {math_result['n_ok']} &nbsp; 🔴 {math_result['n_kritikal']} &nbsp; 🟡 {math_result['n_minor']}</span>
</div>""", unsafe_allow_html=True)

    if n_comp > 0:
        with st.expander(f"📋 Detail Komponen Resume ({n_comp} item)", expanded=True):
            rows_html = ""
            for kode, comp in ext["components"].items():
                usd_str = f"{comp['usd']:,.0f}" if comp.get("usd") else "-"
                exp_usd = ""
                if ext.get("kurs_bi") and comp.get("usd"):
                    e = comp["rp"] / ext["kurs_bi"]
                    diff_pct = abs(e - comp["usd"]) / e if e > 0 else 0
                    icon = "✅" if diff_pct <= 0.01 else "⚠️"
                    exp_usd = f'<span style="color:#6b7280;font-size:10px;">{icon} exp:{round(e):,}</span>'
                rows_html += (
                    f'<tr><td style="padding:6px 10px;font-weight:700;color:#1e6fbf;">{kode}</td>'
                    f'<td style="padding:6px 10px;">{comp["nama"]}</td>'
                    f'<td style="padding:6px 10px;text-align:right;font-family:monospace;">'
                    f'Rp {comp["rp"]:,.0f}</td>'
                    f'<td style="padding:6px 10px;text-align:right;font-family:monospace;">'
                    f'USD {usd_str} {exp_usd}</td></tr>'
                )
            if ext.get("total_rp"):
                rows_html += (
                    f'<tr style="background:#f0f9f4;font-weight:800;border-top:2px solid #1a9e67;">'
                    f'<td colspan="2" style="padding:8px 10px;">TOTAL</td>'
                    f'<td style="padding:8px 10px;text-align:right;font-family:monospace;color:#1a9e67;">'
                    f'Rp {ext["total_rp"]:,.0f}</td>'
                    f'<td style="padding:8px 10px;text-align:right;font-family:monospace;color:#1a9e67;">'
                    f'USD {ext["total_usd"]:,.0f if ext.get("total_usd") else "-"}</td></tr>'
                )
            st.markdown(f"""
<table style="width:100%;border-collapse:collapse;font-size:13px;
              background:#fff;border:1px solid #dde3ea;border-radius:8px;overflow:hidden;">
<thead><tr style="background:#f8fafb;border-bottom:2px solid #dde3ea;">
<th style="padding:8px 10px;text-align:left;font-size:11px;color:#6b7280;width:40px;">Kode</th>
<th style="padding:8px 10px;text-align:left;font-size:11px;color:#6b7280;">Uraian</th>
<th style="padding:8px 10px;text-align:right;font-size:11px;color:#6b7280;">Nilai Rp</th>
<th style="padding:8px 10px;text-align:right;font-size:11px;color:#6b7280;">Nilai USD</th>
</tr></thead>
<tbody>{rows_html}</tbody>
</table>""", unsafe_allow_html=True)

    math_grouped = {}
    for f in mf:
        math_grouped.setdefault(f["severity"], []).append(f)
    for sev in ["kritikal", "minor", "ok", "info"]:
        grp = math_grouped.get(sev, [])
        if not grp: continue
        cfg = SEVERITY_CONFIG[sev]
        st.markdown(
            f'<div style="font-size:11px;font-family:monospace;color:#6b7280;'
            f'text-transform:uppercase;letter-spacing:1.5px;margin:12px 0 8px;">'
            f'{cfg["emoji"]} {sev.upper()} ({len(grp)})'
            f'<span style="display:inline-block;height:1px;background:#dde3ea;'
            f'width:160px;margin-left:10px;vertical-align:middle;"></span></div>',
            unsafe_allow_html=True
        )
        for f in grp:
            formula_html = ""
            if f.get("formula") and f["formula"] != "-":
                formula_html = (
                    f'<div style="margin-top:6px;font-size:11px;background:#f0f4ff;'
                    f'border:1px solid #bfcfee;border-radius:4px;padding:5px 10px;'
                    f'color:#1e6fbf;font-family:monospace;">📐 {f["formula"]}</div>'
                )
            sc = SEVERITY_CONFIG.get(f["severity"], SEVERITY_CONFIG["info"])
            st.markdown(f"""
<div style="background:{sc['bg']};border:1px solid {sc['color']}40;
            border-left:4px solid {sc['color']};border-radius:8px;padding:12px 16px;margin-bottom:8px;">
    <div style="display:flex;align-items:center;gap:8px;margin-bottom:6px;">
        <span style="background:{sc['color']}22;color:{sc['color']};border:1px solid {sc['color']};
                     font-size:10px;font-weight:700;padding:2px 8px;border-radius:4px;
                     text-transform:uppercase;">{sc['emoji']} {f['severity']}</span>
        <span style="color:#666;font-size:11px;font-family:monospace;">{f.get('category','')}</span>
        <span style="color:#666;font-size:11px;font-family:monospace;margin-left:auto;">
            {f.get('id','')} &nbsp; {f.get('page_hint','')}</span>
    </div>
    <div style="font-size:14px;font-weight:600;color:#1a1a2e;margin-bottom:6px;">{f['title']}</div>
    <div style="font-size:12px;color:#444;font-family:monospace;white-space:pre-line;
                background:#ffffff;padding:8px 12px;border-radius:6px;
                border:1px solid #e0e0e0;">{f['detail']}</div>
    {formula_html}
</div>""", unsafe_allow_html=True)


# ══════════════════════════════════════════════
# MAIN APP
# ══════════════════════════════════════════════

def main():
    st.set_page_config(
        page_title="CekLaporan v7 — KJPP SRR",
        page_icon="📋",
        layout="wide",
        initial_sidebar_state="expanded",
    )

    st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Sora:wght@400;600;700;800&family=DM+Mono:wght@400;500&display=swap');
html, body, [class*="css"] { font-family: 'Sora', sans-serif; }
.stApp { background: #f4f6f9; color: #1a1a2e; }
section[data-testid="stSidebar"] { background: #ffffff; border-right: 1px solid #dde3ea; }
section[data-testid="stSidebar"] * { color: #1a1a2e !important; }
.block-container { padding-top: 2rem; max-width: 1100px; }
div[data-testid="stFileUploader"] { border: 2px dashed #c9d2dc; border-radius: 10px; padding: 10px; background: #fff; }
.stButton > button {
    background: #1a9e67; color: #fff; font-weight: 800;
    border: none; border-radius: 8px; padding: 10px 24px;
    font-family: 'Sora', sans-serif; transition: all .2s;
}
.stButton > button:hover { background: #147a50; transform: translateY(-1px); }
.stButton > button:disabled { background: #e5e7eb; color: #9ca3af; }
.stTextInput input { background: #fff; border: 1px solid #dde3ea; border-radius: 6px; font-family: 'DM Mono', monospace; }
.stSelectbox > div > div { background: #fff; border-color: #dde3ea; }
hr { border-color: #dde3ea; }
h1,h2,h3 { color: #1a1a2e !important; font-family: 'Sora', sans-serif !important; }
.stTabs [data-baseweb="tab-list"] { background: #fff; border-bottom: 2px solid #e5e7eb; border-radius: 8px 8px 0 0; }
.stTabs [data-baseweb="tab"] { color: #6b7280; font-size: 13px; font-weight: 600; }
.stTabs [aria-selected="true"] { color: #1a9e67 !important; border-bottom-color: #1a9e67 !important; }
div[data-testid="stExpander"] { background: #fff; border: 1px solid #dde3ea; border-radius: 8px; }
</style>""", unsafe_allow_html=True)

    # ── HEADER ──────────────────────────────────────────────────────────────
    st.markdown("""
<div style="display:flex;align-items:center;gap:12px;margin-bottom:8px;">
    <div style="background:#1a9e67;border-radius:10px;width:40px;height:40px;
                display:flex;align-items:center;justify-content:center;font-size:20px;">📋</div>
    <div>
        <h1 style="margin:0;font-size:24px;font-weight:800;letter-spacing:-0.5px;color:#1a1a2e;">
            Cek<span style="color:#1a9e67;">Laporan</span>
            <span style="font-size:12px;background:#edfaf4;color:#1a9e67;border:1px solid #1a9e67;
                         padding:2px 10px;border-radius:20px;font-weight:600;margin-left:8px;">v7.0 — AI + Excel</span>
        </h1>
        <p style="margin:0;color:#6b7280;font-size:12px;font-family:'DM Mono',monospace;">
            KJPP Suwendho Rinaldy dan Rekan · Pengecekan Laporan Penilaian</p>
    </div>
</div>
<hr style="border-color:#dde3ea;">""", unsafe_allow_html=True)

    # ── KONEKSI GOOGLE SHEETS ────────────────────────────────────────────────
    _, spreadsheet, gs_error = get_gsheet_client()
    gs_connected = spreadsheet is not None
    data_laporan = load_data_laporan(spreadsheet)
    riwayat      = load_riwayat(spreadsheet)

    # ════════════════════════════════════════════════
    # SIDEBAR
    # ════════════════════════════════════════════════
    with st.sidebar:
        if gs_connected:
            st.markdown("""
<div style="background:#edfaf4;border:1px solid #b2e8d0;border-radius:8px;
            padding:10px 12px;margin-bottom:12px;font-size:12px;">
    🟢 <strong>Google Sheets</strong> terhubung
</div>""", unsafe_allow_html=True)
        else:
            st.markdown(f"""
<div style="background:#fff8e6;border:1px solid #f5e0a0;border-radius:8px;
            padding:10px 12px;margin-bottom:12px;font-size:12px;">
    🟡 <strong>Google Sheets</strong> tidak terkonfigurasi<br>
    <span style="color:#6b7280;font-size:11px;">Data hanya tersimpan sementara</span>
</div>""", unsafe_allow_html=True)
            with st.expander("🔍 Detail error", expanded=False):
                st.code(gs_error or "Secrets belum diisi", language="text")

        st.markdown("### 🔑 Claude API Key")
        api_key = st.text_input(
            "Masukkan API Key", type="password",
            placeholder="sk-ant-api03-...",
            help="Dapatkan di https://console.anthropic.com",
        )
        if api_key:
            if api_key.startswith("sk-ant"):
                st.success("✅ Format key valid")
            else:
                st.error("❌ Harus diawali sk-ant")

        st.markdown("---")
        st.markdown("### ⚡ Mode Pengecekan")
        mode_label = st.radio("Pilih mode", options=list(MODE_CONFIG.keys()), label_visibility="collapsed")
        st.caption(MODE_CONFIG[mode_label]["desc"])

        st.markdown("---")
        st.markdown("### ✅ Item yang Dicek")
        _mode_key = MODE_CONFIG[mode_label]["key"]
        if _mode_key == "saham":
            _items_pool = CHECK_ITEMS_SAHAM
            _unchecked  = []
        elif _mode_key == "fairness":
            _items_pool = CHECK_ITEMS_FAIRNESS
            _unchecked  = []
        elif _mode_key == "aset":
            _items_pool = CHECK_ITEMS_ASET
            _unchecked  = ["Konsistensi nomor IMB/perizinan jika ada"] if "Konsistensi nomor IMB/perizinan jika ada" in CHECK_ITEMS_ASET else []
        else:
            _items_pool = CHECK_ITEMS_DEFAULT
            _unchecked  = ["Analisis Pasar & Data Pembanding"]

        selected_items = []
        for item in _items_pool:
            default_checked = item not in _unchecked
            if st.checkbox(item, value=default_checked, key=f"chk_{item}"):
                selected_items.append(item)

        st.markdown("---")
        st.markdown("### 📚 Referensi Laporan")
        laporan_names = list(data_laporan.keys())
        selected_ref  = st.selectbox("Pilih referensi", ["(Tidak ada)"] + laporan_names + ["+ Tambah Baru"])
        if selected_ref not in ["(Tidak ada)", "+ Tambah Baru"] and selected_ref in data_laporan:
            with st.expander("Lihat referensi"):
                for k, v in data_laporan[selected_ref].items():
                    st.write(f"• {k}: **{v}x**")
        if selected_ref == "+ Tambah Baru":
            new_name = st.text_input("Nama laporan baru:")
            if st.button("Tambahkan") and new_name:
                if new_name not in data_laporan:
                    add_laporan_baru(spreadsheet, new_name)
                    st.success(f"✅ '{new_name}' ditambahkan")
                    st.rerun()
                else:
                    st.warning("Nama sudah ada.")

    # ════════════════════════════════════════════════
    # TABS
    # ════════════════════════════════════════════════
    tab_audit, tab_search, tab_history, tab_ref = st.tabs([
        "🤖 AI Audit", "🔍 Pencarian Teks", "📜 Riwayat Audit", "📁 Kelola Referensi"
    ])

    # ────────────────────────────────────────────────
    # TAB 1: AI AUDIT
    # ────────────────────────────────────────────────
    with tab_audit:

        # --- Upload Section ---
        col_doc, col_xl = st.columns([3, 2])

        with col_doc:
            st.markdown("#### 📄 Upload Laporan (PDF/DOCX)")
            uploaded_pdfs  = st.file_uploader("File PDF",  type="pdf",
                accept_multiple_files=True, key="pdf_audit")
            uploaded_docxs = st.file_uploader("File DOCX", type="docx",
                accept_multiple_files=True, key="docx_audit")
            all_files = (uploaded_pdfs or []) + (uploaded_docxs or [])
            if all_files:
                for f in all_files:
                    icon = "📄" if f.name.endswith(".pdf") else "📝"
                    st.markdown(
                        f'<span style="font-family:monospace;font-size:12px;color:#6b7280;">'
                        f'{icon} {f.name} · {f.size//1024} KB</span>',
                        unsafe_allow_html=True
                    )

        with col_xl:
            st.markdown("#### 📊 Upload Lembar Kerja (XLSX) — Opsional")
            if not EXCEL_AVAILABLE:
                st.warning("⚠️ openpyxl tidak tersedia. Install dengan: `pip install openpyxl`")
                uploaded_xlsx = None
            else:
                uploaded_xlsx = st.file_uploader(
                    "File XLSX (lembar kerja/workbook)",
                    type=["xlsx", "xls"],
                    accept_multiple_files=False,
                    key="xlsx_audit",
                    help="Jika diupload, angka di laporan akan dibandingkan dengan data Excel"
                )
                if uploaded_xlsx:
                    st.markdown(
                        f'<span style="font-family:monospace;font-size:12px;color:#1a9e67;">'
                        f'📊 {uploaded_xlsx.name} · {uploaded_xlsx.size//1024} KB · '
                        f'✅ Akan dikomparasi dengan laporan</span>',
                        unsafe_allow_html=True
                    )
                else:
                    st.info("💡 Upload lembar kerja XLSX untuk komparasi angka laporan vs workbook")

        st.markdown("---")
        col_run, col_info = st.columns([2, 5])
        with col_run:
            run_disabled = not (api_key and api_key.startswith("sk-ant") and all_files and selected_items)
            run_btn = st.button("▶ Jalankan Analisis", disabled=run_disabled, use_container_width=True)
        with col_info:
            if not api_key:
                st.info("💡 Masukkan API Key di sidebar.")
            elif not all_files:
                st.info("💡 Upload minimal satu file laporan (PDF atau DOCX).")
            elif not selected_items:
                st.info("💡 Pilih minimal satu item pengecekan di sidebar.")
            else:
                excel_badge = f" · 📊 Excel: {uploaded_xlsx.name}" if uploaded_xlsx else ""
                st.success(f"✅ Siap: {len(all_files)} file · {len(selected_items)} item · **{mode_label}**{excel_badge}")

        if run_btn:
            st.markdown("---")

            # ── Step 1: Baca dokumen ──────────────────────────────────────────
            with st.status("📖 Membaca dokumen...", expanded=True) as status:
                all_pages, file_info = [], []
                for f in all_files:
                    st.write(f"Membaca **{f.name}**...")
                    pages = extract_text_pdf(f) if f.name.endswith(".pdf") else extract_text_docx(f)
                    all_pages.extend(pages)
                    file_info.append(f"{f.name} ({len(pages)} hal.)")
                doc_text = pages_to_text(all_pages)
                st.session_state["doc_pages"] = all_pages
                status.update(label=f"✅ {len(all_pages)} halaman dibaca dari {len(all_files)} file", state="complete")

            # ── Step 2: Parse Excel (jika ada) ───────────────────────────────
            excel_data   = None
            excel_context = ""
            comparison   = {}
            if uploaded_xlsx:
                with st.status("📊 Membaca lembar kerja Excel...", expanded=True) as status:
                    st.write(f"Membaca **{uploaded_xlsx.name}**...")
                    excel_data = parse_excel_workbook(uploaded_xlsx)
                    if excel_data.get("error"):
                        st.warning(f"⚠️ Gagal membaca Excel: {excel_data['error']}")
                        excel_data = None
                    else:
                        n_sheets  = len(excel_data.get("sheets", []))
                        n_numbers = len(excel_data.get("all_numbers", []))
                        status.update(
                            label=f"✅ Excel dibaca: {n_sheets} sheet, {n_numbers} nilai numerik",
                            state="complete"
                        )
                        st.write(excel_data.get("summary", ""))
                        excel_context = build_excel_context_for_prompt(excel_data)

            # ── Step 3: Cek lokal (artefak, placeholder, penomoran) ──────────
            with st.status("🔍 Pengecekan lokal (artefak & placeholder)...", expanded=False) as status:
                local_result = cek_artefak_dan_placeholder(doc_text, _mode_key)
                status.update(
                    label=f"✅ Lokal: {local_result['n_kritikal']}🔴 {local_result['n_minor']}🟡",
                    state="complete"
                )

            # ── Step 4: Cek matematika ───────────────────────────────────────
            math_result = cek_matematika_resume(doc_text)

            # ── Step 5: Komparasi Excel vs dokumen (lokal) ───────────────────
            if excel_data:
                with st.status("🔢 Komparasi angka dokumen vs Excel...", expanded=False) as status:
                    doc_numbers = extract_doc_numbers(doc_text)
                    comparison  = compare_doc_with_excel(doc_numbers, excel_data)
                    n_miss      = len(comparison.get("not_found", []))
                    n_match     = len(comparison.get("matches", []))
                    status.update(
                        label=f"✅ Komparasi: {n_match} cocok, {n_miss} tidak ditemukan di Excel",
                        state="complete"
                    )

            # ── Step 6: Claude AI analisis ────────────────────────────────────
            with st.status("🧠 Claude menganalisis laporan...", expanded=True) as status:
                st.write(f"Model: `{MODEL}` · Mode: **{mode_label}**")
                st.write(f"Teks: **{len(doc_text):,}** karakter")
                if excel_context:
                    st.write(f"Excel context: **{len(excel_context):,}** karakter")
                t_start = time.time()
                try:
                    result, raw_text = call_claude(
                        api_key,
                        MODE_CONFIG[mode_label]["instruction"],
                        selected_items,
                        doc_text,
                        excel_context,
                    )
                    elapsed = time.time() - t_start
                    status.update(label=f"✅ Analisis selesai dalam {elapsed:.1f}s", state="complete")
                except Exception as e:
                    status.update(label="❌ Error", state="error")
                    st.error(f"**Error:** {e}")
                    st.stop()

            # ── Simpan riwayat ────────────────────────────────────────────────
            audit_id = datetime.now().strftime("%Y%m%d_%H%M%S")
            save_riwayat_row(spreadsheet, audit_id, {
                "timestamp": datetime.now().strftime("%d %b %Y, %H:%M"),
                "files":     [f.name for f in all_files],
                "mode":      mode_label,
                "result":    result,
            })
            if gs_connected:
                st.toast("💾 Riwayat tersimpan ke Google Sheets", icon="✅")

            # ════════════════════════════════════════════════
            # TAMPILKAN HASIL
            # ════════════════════════════════════════════════
            st.markdown("---")

            # Header hasil
            report_type = result.get("report_type", "").upper()
            st.markdown(f"### 📊 Hasil Audit — `{report_type}`")
            st.caption(
                f"{' · '.join(file_info)} · "
                f"{datetime.now().strftime('%d %b %Y, %H:%M')} · "
                f"Mode: {mode_label}"
            )

            if result.get("_partial"):
                st.warning(
                    "⚠️ **Response Claude terpotong** karena laporan terlalu panjang. "
                    "Temuan yang berhasil dibaca tetap ditampilkan. "
                    "Coba mode **Pre-Check** untuk laporan besar."
                )

            # Summary cards — gabung AI + lokal
            summary = result.get("summary", {})
            # Support both new (kritis/mayor) and old (kritikal) field names
            ai_kritis = summary.get("kritis", 0) + summary.get("kritikal", 0)
            ai_mayor  = summary.get("mayor", 0)
            ai_minor  = summary.get("minor", 0)
            total_kritis = ai_kritis + local_result["n_kritikal"]
            total_minor  = ai_minor  + local_result["n_minor"]
            merged_summary = {
                **summary,
                "kritis":  total_kritis,
                "mayor":   ai_mayor,
                "minor":   total_minor,
                "overall_score": max(0, summary.get("overall_score", 80)
                                     - local_result["n_kritikal"] * 8
                                     - local_result["n_minor"] * 2),
            }
            render_summary_cards(merged_summary)
            st.markdown("<br>", unsafe_allow_html=True)

            # Executive summary
            exec_sum = summary.get("executive_summary", "")
            if exec_sum:
                st.markdown(f"""
<div style="background:#fff;border:1px solid #dde3ea;border-radius:10px;padding:16px;
            margin-bottom:16px;box-shadow:0 1px 4px rgba(0,0,0,0.05);">
    <div style="font-size:10px;color:#6b7280;font-family:monospace;text-transform:uppercase;
                letter-spacing:1.5px;margin-bottom:8px;">RINGKASAN EKSEKUTIF</div>
    <p style="font-size:14px;line-height:1.8;color:#374151;margin:0;">{exec_sum}</p>
</div>""", unsafe_allow_html=True)

            # Properties
            properties = result.get("properties", [])
            if len(properties) > 1:
                st.markdown(f"""
<div style="background:#eef4ff;border:1px solid #bfcfee;border-radius:10px;
            padding:14px;margin-bottom:16px;">
    <div style="font-size:10px;color:#1e6fbf;font-family:monospace;text-transform:uppercase;
                letter-spacing:1.5px;margin-bottom:8px;">🏢 OBJEK TERDETEKSI ({len(properties)})</div>
    <div style="display:flex;flex-wrap:wrap;gap:6px;">
        {"".join(f'<span style="background:#fff;color:#1e6fbf;border:1px solid #bfcfee;'
                 f'font-size:11px;font-family:monospace;padding:3px 10px;border-radius:12px;">{p}</span>'
                 for p in properties)}
    </div>
</div>""", unsafe_allow_html=True)

            # ── Sub-tabs hasil ──────────────────────────────────────────────
            rtab1, rtab2, rtab3, rtab4 = st.tabs([
                "🔍 Temuan AI",
                "🔎 Cek Lokal (Artefak & Placeholder)",
                "🔢 Matematika",
                "📊 Komparasi Excel",
            ])

            # ── Sub-tab 1: Temuan AI ──────────────────────────────────────
            with rtab1:
                findings = result.get("findings", [])
                filter_prop = None
                if len(properties) > 1:
                    sel = st.selectbox("Filter per objek:", ["Semua Objek"] + properties, key="prop_filter")
                    if sel != "Semua Objek":
                        filter_prop = sel
                filtered_findings = [
                    f for f in findings
                    if not (filter_prop and f.get("property") and f["property"] != filter_prop)
                ]
                _pages = st.session_state.get("doc_pages", [])
                render_findings_as_sections(filtered_findings, _pages)

                with st.expander("🔧 Raw JSON Output (debug)"):
                    st.code(raw_text, language="json")

            # ── Sub-tab 2: Cek Lokal ──────────────────────────────────────
            with rtab2:
                st.markdown("""
<div style="background:#fff8e6;border:1px solid #f5e0a0;border-radius:8px;
            padding:10px 14px;margin-bottom:12px;font-size:12px;color:#7c5800;">
    ⚡ Pengecekan ini berjalan secara lokal (tanpa AI) sehingga hasilnya instan dan deterministik.
    Mencakup: placeholder belum diisi, penomoran ganda, inkonsistensi ejaan.
</div>""", unsafe_allow_html=True)
                lf = local_result.get("findings", [])
                for sev in ["kritikal", "minor", "ok", "info"]:
                    group = [f for f in lf if f.get("severity") == sev]
                    if not group: continue
                    cfg = SEVERITY_CONFIG[sev]
                    st.markdown(
                        f'<div style="font-size:11px;font-family:monospace;color:#6b7280;'
                        f'text-transform:uppercase;letter-spacing:1.5px;margin:12px 0 8px;">'
                        f'{cfg["emoji"]} {sev.upper()} ({len(group)})'
                        f'<span style="display:inline-block;height:1px;background:#dde3ea;'
                        f'width:160px;margin-left:10px;vertical-align:middle;"></span></div>',
                        unsafe_allow_html=True
                    )
                    for f in group:
                        render_finding_card(f)

            # ── Sub-tab 3: Matematika ─────────────────────────────────────
            with rtab3:
                st.markdown("""
<div style="display:flex;align-items:center;gap:8px;margin-bottom:12px;">
    <span style="font-size:18px;">🔢</span>
    <span style="font-size:16px;font-weight:800;color:#1a1a2e;">Pengecekan Matematika Otomatis</span>
    <span style="font-size:11px;background:#edfaf4;color:#1a9e67;border:1px solid #1a9e67;
                 padding:2px 8px;border-radius:12px;font-family:monospace;">Exact Calculation</span>
</div>""", unsafe_allow_html=True)
                render_math_section(math_result)

            # ── Sub-tab 4: Komparasi Excel ────────────────────────────────
            with rtab4:
                if not uploaded_xlsx:
                    st.info(
                        "📊 **Tidak ada lembar kerja yang diupload.**\n\n"
                        "Upload file XLSX di atas bersamaan dengan PDF/DOCX laporan untuk "
                        "membandingkan angka-angka di laporan dengan data di lembar kerja. "
                        "Ini membantu mendeteksi:\n"
                        "- Angka di laporan yang berbeda dari lembar kerja\n"
                        "- Angka di laporan yang tidak ada di lembar kerja\n"
                        "- Kesalahan pembulatan atau konversi"
                    )
                elif excel_data and excel_data.get("error"):
                    st.error(f"❌ Gagal membaca Excel: {excel_data['error']}")
                elif comparison:
                    st.markdown(f"""
<div style="background:#f5f3ff;border:1px solid #c4b5fd;border-radius:10px;
            padding:14px 16px;margin-bottom:12px;">
    <div style="font-size:10px;color:#7c3aed;font-family:monospace;text-transform:uppercase;
                letter-spacing:1.5px;margin-bottom:8px;">📊 KOMPARASI LEMBAR KERJA VS LAPORAN</div>
    <p style="margin:0;font-size:13px;color:#374151;">{comparison.get('summary','')}</p>
</div>""", unsafe_allow_html=True)

                    # Info Excel
                    st.markdown("**Sheet yang dibaca:**")
                    for sheet in excel_data.get("sheets", []):
                        n = len(sheet["rows"])
                        if n > 0:
                            st.markdown(
                                f'<span style="font-family:monospace;font-size:12px;'
                                f'color:#6b7280;">📋 {sheet["name"]}: {n} nilai numerik</span>',
                                unsafe_allow_html=True
                            )
                    st.markdown("---")
                    render_excel_comparison(comparison)
                else:
                    st.info("Tidak ada data komparasi.")

            # ── Download Report ──────────────────────────────────────────────
            st.markdown("---")
            report_text = generate_download_report(
                result, math_result, local_result,
                [f.name for f in all_files], mode_label
            )
            st.download_button(
                label="⬇️ Download Laporan Review (.txt)",
                data=report_text.encode("utf-8"),
                file_name=f"review_{audit_id}.txt",
                mime="text/plain",
                use_container_width=False,
            )

    # ────────────────────────────────────────────────
    # TAB 2: PENCARIAN TEKS
    # ────────────────────────────────────────────────
    with tab_search:
        st.markdown("#### 🔍 Pencarian Frasa Manual")
        st.caption("Upload dokumen dan cari frasa — dilengkapi highlight.")
        col_s1, col_s2 = st.columns([1, 2])
        with col_s1:
            up_pdf_s  = st.file_uploader("PDF",  type="pdf",  accept_multiple_files=True, key="pdf_search")
            up_docx_s = st.file_uploader("DOCX", type="docx", accept_multiple_files=True, key="docx_search")
            files_s   = (up_pdf_s or []) + (up_docx_s or [])
            sel_ref   = st.selectbox("Laporan Referensi", ["(Tidak ada)"] + list(data_laporan.keys()), key="_sel_laporan_s2")
            if sel_ref != "(Tidak ada)" and sel_ref in data_laporan:
                st.markdown("**Wajib cek:**")
                for k, v in data_laporan[sel_ref].items():
                    st.write(f"• {k}: {v}x")
        with col_s2:
            phrase = st.text_input("Frasa yang dicari:", placeholder="contoh: tanggal penilaian")
            if phrase and files_s:
                grand_total = 0
                for uf in files_s:
                    pages = extract_text_pdf(uf) if uf.name.endswith(".pdf") else extract_text_docx(uf)
                    pattern    = re.compile(r"\s*".join(re.escape(w) for w in phrase.split()), re.IGNORECASE)
                    file_total = 0
                    st.markdown(f"**📄 {uf.name}**")
                    found_any  = False
                    for i, page_txt in enumerate(pages, 1):
                        matches = pattern.findall(page_txt)
                        if matches:
                            found_any   = True
                            file_total += len(matches)
                            highlighted = re.sub(
                                fr"({re.escape(phrase)})",
                                r'<mark style="background:#fbbf24;color:#000;padding:0 2px;">\1</mark>',
                                page_txt, flags=re.IGNORECASE
                            )
                            with st.expander(f"Halaman {i} — {len(matches)} kemunculan"):
                                st.markdown(highlighted, unsafe_allow_html=True)
                    if not found_any:
                        st.caption("Tidak ditemukan.")
                    else:
                        st.markdown(f"→ **Total di file ini: {file_total}x**")
                    grand_total += file_total
                st.markdown(f"---\n### Total semua file: **{grand_total}x**")
            elif phrase and not files_s:
                st.info("Upload file terlebih dahulu.")

    # ────────────────────────────────────────────────
    # TAB 3: RIWAYAT AUDIT
    # ────────────────────────────────────────────────
    with tab_history:
        st.markdown("#### 📜 Riwayat Audit")
        if gs_connected:
            st.markdown('<span style="background:#edfaf4;color:#1a9e67;border:1px solid #1a9e67;'
                        'font-size:11px;padding:2px 10px;border-radius:12px;">💾 Google Sheets</span>',
                        unsafe_allow_html=True)
        else:
            st.markdown('<span style="background:#fff8e6;color:#d4860a;border:1px solid #d4860a;'
                        'font-size:11px;padding:2px 10px;border-radius:12px;">⚠️ Sementara (session)</span>',
                        unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)

        riwayat = load_riwayat(spreadsheet)
        if not riwayat:
            st.info("Belum ada riwayat audit.")
        else:
            st.caption(f"Total: **{len(riwayat)}** audit tersimpan")
            for aid in sorted(riwayat.keys(), reverse=True):
                r      = riwayat[aid]
                score  = r.get("score", 0)
                fstr   = ", ".join(r.get("files", []))
                with st.expander(f"🕒 {r.get('timestamp','')}  ·  {fstr}  ·  Skor: {score}"):
                    c1, c2, c3, c4 = st.columns(4)
                    c1.metric("Skor QC",    score)
                    c2.metric("🔴 Kritikal", r.get("kritikal", 0))
                    c3.metric("🟡 Minor",    r.get("minor", 0))
                    c4.metric("Mode",        r.get("mode", ""))
                    exec_s = r.get("result", {}).get("summary", {}).get("executive_summary", "")
                    if exec_s:
                        st.caption(exec_s)
                    if st.button("Tampilkan Detail", key=f"hist_{aid}"):
                        for f in r.get("result", {}).get("findings", []):
                            render_finding_card(f)
            st.markdown("---")
            if st.button("🗑 Hapus Semua Riwayat"):
                clear_riwayat(spreadsheet)
                st.success("Riwayat dihapus.")
                st.rerun()

    # ────────────────────────────────────────────────
    # TAB 4: KELOLA REFERENSI
    # ────────────────────────────────────────────────
    with tab_ref:
        st.markdown("#### 📁 Kelola Data Referensi Laporan")
        st.caption("Simpan catatan frekuensi kemunculan frasa per jenis laporan sebagai baseline.")
        if gs_connected:
            st.markdown('<span style="background:#edfaf4;color:#1a9e67;border:1px solid #1a9e67;'
                        'font-size:11px;padding:2px 10px;border-radius:12px;">💾 Google Sheets</span>',
                        unsafe_allow_html=True)
            st.markdown("<br>", unsafe_allow_html=True)

        data_laporan = load_data_laporan(spreadsheet)
        if not data_laporan:
            st.info("Belum ada data referensi.")
        else:
            for lap_name, lap_data in data_laporan.items():
                with st.expander(f"📋 {lap_name}"):
                    if lap_data:
                        for k, v in lap_data.items():
                            c1, c2, c3 = st.columns([3, 1, 1])
                            c1.write(k)
                            c2.write(f"**{v}x**")
                            if c3.button("🗑", key=f"del_{lap_name}_{k}"):
                                delete_referensi_row(spreadsheet, lap_name, k)
                                st.rerun()
                    else:
                        st.caption("Belum ada keterangan.")
                    st.markdown("**Tambah keterangan:**")
                    ck, cv, cbtn = st.columns([3, 1, 1])
                    new_k = ck.text_input("Frasa",  key=f"nk_{lap_name}")
                    new_v = cv.number_input("Jumlah", min_value=0, step=1, key=f"nv_{lap_name}")
                    if cbtn.button("Tambah", key=f"nbtn_{lap_name}") and new_k:
                        save_referensi_row(spreadsheet, lap_name, new_k, new_v)
                        st.success("✅ Ditambahkan")
                        st.rerun()

        st.markdown("---")
        st.markdown("**Tambah Laporan Referensi Baru:**")
        c1, c2 = st.columns([3, 1])
        new_lap = c1.text_input("Nama laporan:", key="new_lap_name")
        if c2.button("Tambahkan", key="btn_new_lap") and new_lap:
            if new_lap not in data_laporan:
                add_laporan_baru(spreadsheet, new_lap)
                st.success(f"✅ '{new_lap}' ditambahkan")
                st.rerun()
            else:
                st.warning("Nama sudah ada.")

    # ── FOOTER ──────────────────────────────────────────────────────────────
    st.markdown("""
<hr style="margin-top:40px;border-color:#dde3ea;">
<p style="text-align:center;font-size:12px;color:#9ca3af;font-family:'DM Mono',monospace;">
    CekLaporan v7.0 &nbsp;·&nbsp; KJPP SRR &nbsp;·&nbsp;
    Powered by Claude AI (claude-sonnet-4-5) &nbsp;·&nbsp; Created by HW
</p>""", unsafe_allow_html=True)


if __name__ == "__main__":
    main()
